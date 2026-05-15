from flask import Flask, jsonify, request, render_template, send_from_directory, session
from flask_cors import CORS
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from werkzeug.utils import secure_filename
import sqlite3
import json
import os
import uuid
import secrets
import requests
import bcrypt
from datetime import datetime
from dotenv import load_dotenv
load_dotenv()

app = Flask(__name__)
CORS(app)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB
app.secret_key = os.environ.get('SECRET_KEY', 'dev-fallback-key-change-me')

# 管理员密码
ADMIN_PASSWORD_HASH = bcrypt.hashpw(os.environ.get('ADMIN_PASSWORD', 'changeme').encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

# ========== 限流防护 ==========
limiter = Limiter(
    get_remote_address,
    app=app,
    default_limits=["60 per minute"],
    storage_uri="memory://",
)

DB_PATH = os.path.join(os.path.dirname(__file__), 'emotion.db')

# ========== 管理员认证 ==========
from functools import wraps

def admin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get('admin_logged_in'):
            return jsonify({'error': '未授权'}), 401
        return f(*args, **kwargs)
    return decorated

# ========== 文件上传 ==========
UPLOAD_DIR = os.path.join(os.path.dirname(__file__), 'uploads')
os.makedirs(UPLOAD_DIR, exist_ok=True)
ALLOWED_EXT = {'png', 'jpg', 'jpeg', 'gif', 'webp', 'mp4', 'mov', 'avi', 'mp3', 'wav', 'ogg', 'm4a', 'aac'}

# ========== APP版本管理 ==========
APK_DIR = os.path.join(os.path.dirname(__file__), 'apk')
os.makedirs(APK_DIR, exist_ok=True)
VERSION_FILE = os.path.join(os.path.dirname(__file__), 'version.json')

def get_version_info():
    if os.path.exists(VERSION_FILE):
        with open(VERSION_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {"version_code": 1, "version_name": "1.0", "changelog": "", "apk_name": ""}

@app.route('/api/app/version', methods=['GET'])
def check_app_version():
    return jsonify(get_version_info())

@app.route('/apk/<path:filename>')
def serve_apk(filename):
    return send_from_directory(APK_DIR, filename)

@app.route('/api/app/upload_apk', methods=['POST'])
@limiter.limit("2 per minute")
def upload_apk():
    if 'file' not in request.files:
        return jsonify({'error': '没有文件'}), 400
    f = request.files['file']
    version_code = request.form.get('version_code', '1')
    version_name = request.form.get('version_name', '1.0')
    changelog = request.form.get('changelog', '')
    filename = 'xinyu.apk'
    f.save(os.path.join(APK_DIR, filename))
    # 更新版本信息
    info = {
        "version_code": int(version_code),
        "version_name": version_name,
        "changelog": changelog,
        "apk_name": filename,
        "download_url": "/apk/" + filename
    }
    with open(VERSION_FILE, 'w', encoding='utf-8') as fp:
        json.dump(info, fp, ensure_ascii=False, indent=2)
    return jsonify({'ok': True, 'version_name': version_name})

@app.route('/uploads/<path:filename>')
def serve_upload(filename):
    return send_from_directory(UPLOAD_DIR, filename)

@app.route('/api/upload', methods=['POST'])
@limiter.limit("10 per minute")
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': '没有文件'}), 400
    f = request.files['file']
    if f.filename == '':
        return jsonify({'error': '文件名为空'}), 400
    ext = f.filename.rsplit('.', 1)[-1].lower() if '.' in f.filename else ''
    if ext not in ALLOWED_EXT:
        return jsonify({'error': '不支持的文件格式'}), 400
    filename = str(uuid.uuid4())[:8] + '.' + ext
    f.save(os.path.join(UPLOAD_DIR, filename))
    url = '/uploads/' + filename
    return jsonify({'ok': True, 'url': url})

# ========== 课表导入 ==========
WEEKDAY_MAP = {
    '星期一': 0, '周一': 0, 'monday': 0, 'mon': 0,
    '星期二': 1, '周二': 1, 'tuesday': 1, 'tue': 1,
    '星期三': 2, '周三': 2, 'wednesday': 2, 'wed': 2,
    '星期四': 3, '周四': 3, 'thursday': 3, 'thu': 3,
    '星期五': 4, '周五': 4, 'friday': 4, 'fri': 4,
    '星期六': 5, '周六': 5, 'saturday': 5, 'sat': 5,
    '星期日': 6, '周日': 6, '星期天': 6, '周天': 6, 'sunday': 6, 'sun': 6,
}

def parse_weekday(text):
    """从文本中识别星期几"""
    t = text.lower().strip()
    for key, val in WEEKDAY_MAP.items():
        if key in t:
            return val
    return -1

def parse_time_info(text):
    """从文本中提取节次信息，如 '第1-2节' '1-2节' '1,2节'"""
    import re
    t = text.strip()
    # 匹配 "第X-Y节" 或 "X-Y节" 或 "X,Y节"
    m = re.search(r'(?:第)?(\d+)[\s]*[-,~]\s*(\d+)[\s]*节', t)
    if m:
        return f"第{m.group(1)}-{m.group(2)}节"
    # 匹配 "第X节"
    m = re.search(r'(?:第)?(\d+)[\s]*节', t)
    if m:
        return f"第{m.group(1)}节"
    return ""

def parse_room(text):
    """从文本中提取教室信息"""
    import re
    # 匹配常见教室格式：A301, 教学楼A301, 实验楼B201, 301教室 等
    m = re.search(r'([A-Za-z]?\d{2,4}[A-Za-z]?)\s*(教室|室)?', text)
    if m:
        room = m.group(1)
        if len(room) >= 3:
            return room
    # 匹配 "XX楼XXX"
    m = re.search(r'([\u4e00-\u9fa5]+楼[A-Za-z]?\d{2,4})', text)
    if m:
        return m.group(1)
    return ""

def parse_excel_course(filepath):
    """解析Excel课表文件"""
    import openpyxl
    courses = []
    wb = openpyxl.load_workbook(filepath, data_only=True)
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue

        # 尝试识别表头行和星期列
        header_row = -1
        weekday_cols = {}  # col_index -> weekday

        for ri, row in enumerate(rows):
            if row is None:
                continue
            row_str = [str(c).strip() if c else '' for c in row]
            # 检查是否是表头行（包含星期关键词）
            found_weekdays = 0
            for ci, cell in enumerate(row_str):
                wd = parse_weekday(cell)
                if wd >= 0:
                    weekday_cols[ci] = wd
                    found_weekdays += 1
            if found_weekdays >= 2:
                header_row = ri
                break

        if header_row < 0:
            # 没找到表头，尝试按列解析（第一列是节次，后续列是星期）
            for ri, row in enumerate(rows):
                if row is None:
                    continue
                row_str = [str(c).strip() if c else '' for c in row]
                if not row_str[0]:
                    continue
                time_info = parse_time_info(row_str[0])
                if not time_info:
                    continue
                for ci in range(1, len(row_str)):
                    cell = row_str[ci].strip()
                    if not cell or cell == 'None':
                        continue
                    wd = ci - 1 if ci - 1 < 7 else -1
                    if wd < 0:
                        continue
                    lines = [l.strip() for l in cell.split('\n') if l.strip()]
                    if lines:
                        name = lines[0]
                        room = ""
                        for l in lines[1:]:
                            r = parse_room(l)
                            if r:
                                room = r
                                break
                        courses.append({
                            'name': name,
                            'room': room,
                            'time': time_info,
                            'weekday': wd
                        })
            continue

        # 有表头行，按列解析
        for ri in range(header_row + 1, len(rows)):
            row = rows[ri]
            if row is None:
                continue
            row_str = [str(c).strip() if c else '' for c in row]
            # 第一列通常是节次/时间
            time_info = parse_time_info(row_str[0]) if row_str[0] else ""
            for ci, wd in weekday_cols.items():
                if ci >= len(row_str):
                    continue
                cell = row_str[ci].strip()
                if not cell or cell == 'None':
                    continue
                lines = [l.strip() for l in cell.split('\n') if l.strip()]
                if lines:
                    name = lines[0]
                    room = ""
                    for l in lines[1:]:
                        r = parse_room(l)
                        if r:
                            room = r
                            break
                    t = time_info if time_info else ""
                    courses.append({
                        'name': name,
                        'room': room,
                        'time': t,
                        'weekday': wd
                    })
    wb.close()
    return courses

def parse_pdf_course(filepath):
    """解析PDF课表文件"""
    import pdfplumber
    courses = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()
            for table in tables:
                if not table or len(table) < 2:
                    continue
                # 尝试找到表头行
                header_row = -1
                weekday_cols = {}
                for ri, row in enumerate(table):
                    if not row:
                        continue
                    found = 0
                    for ci, cell in enumerate(row):
                        if cell:
                            wd = parse_weekday(str(cell))
                            if wd >= 0:
                                weekday_cols[ci] = wd
                                found += 1
                    if found >= 2:
                        header_row = ri
                        break

                if header_row >= 0:
                    for ri in range(header_row + 1, len(table)):
                        row = table[ri]
                        if not row:
                            continue
                        time_info = parse_time_info(str(row[0])) if row[0] else ""
                        for ci, wd in weekday_cols.items():
                            if ci >= len(row) or not row[ci]:
                                continue
                            cell = str(row[ci]).strip()
                            if not cell:
                                continue
                            lines = [l.strip() for l in cell.split('\n') if l.strip()]
                            if lines:
                                name = lines[0]
                                room = ""
                                for l in lines[1:]:
                                    r = parse_room(l)
                                    if r:
                                        room = r
                                        break
                                courses.append({
                                    'name': name,
                                    'room': room,
                                    'time': time_info,
                                    'weekday': wd
                                })
                else:
                    # 无表头，尝试逐行解析
                    for row in table:
                        if not row:
                            continue
                        row_str = [str(c).strip() if c else '' for c in row]
                        time_info = parse_time_info(row_str[0]) if row_str[0] else ""
                        if not time_info:
                            continue
                        for ci in range(1, len(row_str)):
                            cell = row_str[ci].strip()
                            if not cell or cell == 'None':
                                continue
                            wd = ci - 1 if ci - 1 < 7 else -1
                            if wd < 0:
                                continue
                            lines = [l.strip() for l in cell.split('\n') if l.strip()]
                            if lines:
                                name = lines[0]
                                room = ""
                                for l in lines[1:]:
                                    r = parse_room(l)
                                    if r:
                                        room = r
                                        break
                                courses.append({
                                    'name': name,
                                    'room': room,
                                    'time': time_info,
                                    'weekday': wd
                                })
    return courses

def parse_word_course(filepath):
    """解析Word课表文件"""
    from docx import Document
    courses = []
    doc = Document(filepath)
    # 先尝试解析表格
    for table in doc.tables:
        rows_data = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_data.append(cells)
        if len(rows_data) < 2:
            continue
        # 找表头
        header_row = -1
        weekday_cols = {}
        for ri, row in enumerate(rows_data):
            found = 0
            for ci, cell in enumerate(row):
                wd = parse_weekday(cell)
                if wd >= 0:
                    weekday_cols[ci] = wd
                    found += 1
            if found >= 2:
                header_row = ri
                break
        if header_row >= 0:
            for ri in range(header_row + 1, len(rows_data)):
                row = rows_data[ri]
                time_info = parse_time_info(row[0]) if row[0] else ""
                for ci, wd in weekday_cols.items():
                    if ci >= len(row) or not row[ci]:
                        continue
                    cell = row[ci].strip()
                    if not cell:
                        continue
                    lines = [l.strip() for l in cell.split('\n') if l.strip()]
                    if lines:
                        name = lines[0]
                        room = ""
                        for l in lines[1:]:
                            r = parse_room(l)
                            if r:
                                room = r
                                break
                        courses.append({
                            'name': name,
                            'room': room,
                            'time': time_info,
                            'weekday': wd
                        })
    # 如果表格没解析出来，尝试从段落文本解析
    if not courses:
        full_text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        lines = full_text.split('\n')
        current_weekday = -1
        for line in lines:
            line = line.strip()
            if not line:
                continue
            wd = parse_weekday(line)
            if wd >= 0:
                current_weekday = wd
                continue
            if current_weekday >= 0:
                time_info = parse_time_info(line)
                if time_info:
                    continue
                # 尝试分割课程名和教室
                parts = line.split()
                if parts:
                    name = parts[0]
                    room = parse_room(line)
                    courses.append({
                        'name': name,
                        'room': room,
                        'time': '',
                        'weekday': current_weekday
                    })
    return courses

@app.route('/api/course/import', methods=['POST'])
@limiter.limit("10 per minute")
def import_course():
    if 'file' not in request.files:
        return jsonify({'error': '没有文件'}), 400
    f = request.files['file']
    if f.filename == '':
        return jsonify({'error': '文件名为空'}), 400
    ext = f.filename.rsplit('.', 1)[-1].lower() if '.' in f.filename else ''
    if ext not in ('pdf', 'doc', 'docx', 'xls', 'xlsx'):
        return jsonify({'error': '不支持的文件格式，请上传PDF、Word或Excel文件'}), 400
    # 保存临时文件
    tmp_path = os.path.join(UPLOAD_DIR, 'tmp_course_' + str(uuid.uuid4())[:8] + '.' + ext)
    f.save(tmp_path)
    try:
        if ext == 'xlsx' or ext == 'xls':
            courses = parse_excel_course(tmp_path)
        elif ext == 'pdf':
            courses = parse_pdf_course(tmp_path)
        elif ext in ('doc', 'docx'):
            courses = parse_word_course(tmp_path)
        else:
            courses = []
        return jsonify({'courses': courses, 'count': len(courses)})
    except Exception as e:
        return jsonify({'error': '解析失败: ' + str(e)}), 500
    finally:
        try:
            os.remove(tmp_path)
        except:
            pass

# ========== 数据库初始化 ==========
def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    conn = get_db()
    conn.executescript('''
        CREATE TABLE IF NOT EXISTS moods (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            mood TEXT NOT NULL,
            note TEXT,
            gender TEXT DEFAULT 'male',
            device_id TEXT DEFAULT '',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        );

        CREATE TABLE IF NOT EXISTS checkins (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            day TEXT NOT NULL,
            completed INTEGER DEFAULT 0,
            gender TEXT DEFAULT 'male',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        );

        CREATE TABLE IF NOT EXISTS chat_history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            role TEXT NOT NULL,
            message TEXT NOT NULL,
            gender TEXT DEFAULT 'male',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        );

        CREATE TABLE IF NOT EXISTS students (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            student_id TEXT UNIQUE NOT NULL,
            name TEXT NOT NULL,
            gender TEXT,
            age INTEGER,
            major TEXT,
            class_name TEXT,
            phone TEXT,
            email TEXT,
            address TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        );

        CREATE TABLE IF NOT EXISTS emergency_contacts (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT NOT NULL,
            name TEXT NOT NULL,
            phone TEXT NOT NULL,
            relation TEXT DEFAULT '',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        );
    ''')
    conn.commit()
    conn.close()

init_db()

# ========== 用户注册/登录（服务器端） ==========
def init_user_table():
    conn = get_db()
    conn.execute('''CREATE TABLE IF NOT EXISTS server_users (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT UNIQUE NOT NULL,
        password TEXT NOT NULL,
        nickname TEXT,
        gender TEXT DEFAULT 'male',
        avatar TEXT DEFAULT '',
        banned INTEGER DEFAULT 0,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')
    # Add banned column if missing (for existing DBs)
    try:
        conn.execute('ALTER TABLE server_users ADD COLUMN banned INTEGER DEFAULT 0')
    except:
        pass
    # Add session tracking columns for single-session enforcement
    try:
        conn.execute('ALTER TABLE server_users ADD COLUMN session_token TEXT')
    except:
        pass
    try:
        conn.execute('ALTER TABLE server_users ADD COLUMN login_platform TEXT')
    except:
        pass
    # Add security question columns for password recovery
    try:
        conn.execute('ALTER TABLE server_users ADD COLUMN security_question TEXT DEFAULT ""')
    except:
        pass
    try:
        conn.execute('ALTER TABLE server_users ADD COLUMN security_answer TEXT DEFAULT ""')
    except:
        pass
    # 新增个人属性字段
    try:
        conn.execute('ALTER TABLE server_users ADD COLUMN age INTEGER DEFAULT 0')
    except:
        pass
    try:
        conn.execute('ALTER TABLE server_users ADD COLUMN role TEXT DEFAULT ""')
    except:
        pass
    try:
        conn.execute('ALTER TABLE server_users ADD COLUMN region TEXT DEFAULT ""')
    except:
        pass
    try:
        conn.execute('ALTER TABLE server_users ADD COLUMN bio TEXT DEFAULT ""')
    except:
        pass
    # 新增用户类型字段
    try:
        conn.execute('ALTER TABLE server_users ADD COLUMN user_type TEXT DEFAULT "user"')
    except:
        pass
    try:
        conn.execute('ALTER TABLE server_users ADD COLUMN verify_status TEXT DEFAULT ""')
    except:
        pass
    try:
        conn.execute('ALTER TABLE server_users ADD COLUMN verify_data TEXT DEFAULT ""')
    except:
        pass
    conn.commit()
    conn.close()

init_user_table()

# ===== 单点登录：token校验装饰器 =====
def login_required(f):
    from functools import wraps
    from urllib.parse import unquote
    @wraps(f)
    def decorated(*args, **kwargs):
        token = request.headers.get('X-Session-Token') or request.args.get('session_token') or (request.json or {}).get('session_token', '')
        username = request.headers.get('X-Username') or request.args.get('username') or (request.json or {}).get('username', '')
        if token: token = unquote(token)
        if username: username = unquote(username)
        if not token or not username:
            return jsonify({'error': '请先登录', 'code': 'SESSION_EXPIRED'}), 401
        conn = get_db()
        user = conn.execute('SELECT session_token FROM server_users WHERE username=?', (username,)).fetchone()
        conn.close()
        if not user or user['session_token'] != token:
            return jsonify({'error': '账号已在其他设备登录', 'code': 'SESSION_EXPIRED'}), 401
        return f(*args, **kwargs)
    return decorated

@app.route('/api/user/register', methods=['POST'])
@limiter.limit("5 per minute")
def register():
    data = request.json or {}
    username = data.get('username', '').strip()
    password = data.get('password', '')
    nickname = data.get('nickname', '').strip()
    gender = data.get('gender', '')
    avatar = data.get('avatar', '')
    user_type = data.get('user_type', 'user')  # user / volunteer / counselor
    if not username or not password:
        return jsonify({'error': '用户名和密码不能为空'}), 400
    if len(password) < 4:
        return jsonify({'error': '密码至少4位'}), 400
    if not gender:
        return jsonify({'error': '请选择性别'}), 400
    if user_type not in ('user', 'volunteer', 'counselor'):
        user_type = 'user'
    # 志愿者和心理辅导师需要额外字段
    verify_data = ''
    verify_status = ''
    if user_type == 'volunteer':
        university = data.get('university', '').strip()
        student_id = data.get('student_id', '').strip()
        reason = data.get('reason', '').strip()
        if not university or not student_id:
            return jsonify({'error': '请填写学校和学号'}), 400
        verify_data = json.dumps({'university': university, 'student_id': student_id, 'reason': reason}, ensure_ascii=False)
        verify_status = 'pending'
    elif user_type == 'counselor':
        real_name = data.get('real_name', '').strip()
        institution = data.get('institution', '').strip()
        license_no = data.get('license_no', '').strip()
        specialization = data.get('specialization', '').strip()
        if not real_name or not institution or not license_no:
            return jsonify({'error': '请填写真实姓名、执业机构和资格证号'}), 400
        verify_data = json.dumps({'real_name': real_name, 'institution': institution, 'license_no': license_no, 'specialization': specialization}, ensure_ascii=False)
        verify_status = 'pending'
    # bcrypt 加密密码
    hashed = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
    platform = data.get('platform', 'web')
    token = secrets.token_hex(16)
    conn = get_db()
    try:
        conn.execute('INSERT INTO server_users (username, password, nickname, gender, avatar, session_token, login_platform, user_type, verify_status, verify_data) VALUES (?,?,?,?,?,?,?,?,?,?)',
                     (username, hashed, nickname, gender, avatar, token, platform, user_type, verify_status, verify_data))
        conn.commit()
        conn.close()
        msg = '注册成功'
        if user_type == 'volunteer':
            msg = '注册成功，志愿者身份需审核通过后生效'
        elif user_type == 'counselor':
            msg = '注册成功，心理辅导师身份需审核通过后生效'
        return jsonify({'ok': True, 'nickname': nickname, 'gender': gender, 'avatar': avatar, 'session_token': token, 'login_platform': platform, 'msg': msg})
    except sqlite3.IntegrityError:
        conn.close()
        return jsonify({'error': '用户名已存在'}), 400

@app.route('/api/user/login', methods=['POST'])
@limiter.limit("5 per minute")
def login():
    data = request.json or {}
    username = data.get('username', '').strip()
    password = data.get('password', '')
    platform = data.get('platform', 'web')
    conn = get_db()
    user = conn.execute('SELECT * FROM server_users WHERE username=?', (username,)).fetchone()
    if user and user['banned']:
        conn.close()
        return jsonify({'error': '账号已被封禁，请联系管理员'}), 403
    if user:
        stored_pwd = user['password']
        pwd_ok = False
        # 兼容旧密码：如果旧密码是明文，验证通过后自动升级为 bcrypt
        if stored_pwd.startswith('$2'):
            pwd_ok = bcrypt.checkpw(password.encode('utf-8'), stored_pwd.encode('utf-8'))
        else:
            if stored_pwd == password:
                pwd_ok = True
                new_hashed = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
                conn.execute('UPDATE server_users SET password=? WHERE username=?', (new_hashed, username))
        if pwd_ok:
            # 生成新token，覆盖旧token（踢掉另一端）
            token = secrets.token_hex(16)
            conn.execute('UPDATE server_users SET session_token=?, login_platform=? WHERE username=?',
                         (token, platform, username))
            conn.commit()
            conn.close()
            return jsonify({'ok': True, 'nickname': user['nickname'], 'gender': user['gender'],
                           'avatar': user['avatar'], 'session_token': token, 'login_platform': platform,
                           'user_type': user['user_type'], 'verify_status': user['verify_status']})
    conn.close()
    return jsonify({'error': '用户名或密码错误'}), 401

@app.route('/api/user/validate', methods=['POST'])
def validate_session():
    """验证session是否有效（用于客户端启动时检查是否被踢下线）"""
    data = request.json or {}
    username = data.get('username', '')
    token = data.get('session_token', '')
    if not username or not token:
        return jsonify({'valid': False})
    conn = get_db()
    user = conn.execute('SELECT session_token, nickname, gender, avatar, user_type, verify_status, age, role, region, bio FROM server_users WHERE username=?', (username,)).fetchone()
    conn.close()
    if user and user['session_token'] == token:
        return jsonify({'valid': True, 'nickname': user['nickname'], 'gender': user['gender'], 'avatar': user['avatar'],
                        'age': user['age'], 'role': user['role'], 'region': user['region'], 'bio': user['bio'],
                        'user_type': user['user_type'], 'verify_status': user['verify_status']})
    return jsonify({'valid': False})

@app.route('/api/user/security/question', methods=['GET'])
def get_security_question():
    """获取安全问题"""
    username = request.args.get('username', '')
    if not username:
        return jsonify({'error': '请输入用户名'}), 400
    conn = get_db()
    user = conn.execute('SELECT security_question FROM server_users WHERE username=?', (username,)).fetchone()
    conn.close()
    if not user:
        return jsonify({'error': '用户不存在'}), 404
    if not user['security_question']:
        return jsonify({'error': '该用户未设置安全问题'}), 400
    return jsonify({'question': user['security_question']})

@app.route('/api/user/security/verify', methods=['POST'])
def verify_security_answer():
    """验证安全答案并返回重置token"""
    data = request.json or {}
    username = data.get('username', '')
    answer = data.get('answer', '').strip()
    if not username or not answer:
        return jsonify({'error': '请填写完整'}), 400
    conn = get_db()
    user = conn.execute('SELECT security_question, security_answer FROM server_users WHERE username=?', (username,)).fetchone()
    if not user:
        conn.close()
        return jsonify({'error': '用户不存在'}), 404
    if not user['security_answer']:
        conn.close()
        return jsonify({'error': '该用户未设置安全问题'}), 400
    if user['security_answer'].lower() != answer.lower():
        conn.close()
        return jsonify({'error': '答案不正确'}), 400
    # 生成重置token，5分钟有效
    reset_token = secrets.token_hex(16)
    import time
    conn.execute('UPDATE server_users SET session_token=? WHERE username=?',
                 (f'reset:{reset_token}:{int(time.time())}', username))
    conn.commit()
    conn.close()
    return jsonify({'ok': True, 'reset_token': reset_token})

@app.route('/api/user/security/reset', methods=['POST'])
def reset_password():
    """用重置token重置密码"""
    data = request.json or {}
    username = data.get('username', '')
    reset_token = data.get('reset_token', '')
    new_password = data.get('new_password', '')
    if not username or not reset_token or not new_password:
        return jsonify({'error': '参数缺失'}), 400
    if len(new_password) < 4:
        return jsonify({'error': '密码至少4位'}), 400
    conn = get_db()
    user = conn.execute('SELECT session_token FROM server_users WHERE username=?', (username,)).fetchone()
    if not user:
        conn.close()
        return jsonify({'error': '用户不存在'}), 404
    # 验证reset token
    token_parts = (user['session_token'] or '').split(':')
    if len(token_parts) != 3 or token_parts[0] != 'reset' or token_parts[1] != reset_token:
        conn.close()
        return jsonify({'error': '重置链接无效'}), 400
    import time
    if int(time.time()) - int(token_parts[2]) > 300:
        conn.close()
        return jsonify({'error': '重置链接已过期，请重新验证'}), 400
    # 重置密码
    new_hashed = bcrypt.hashpw(new_password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
    conn.execute('UPDATE server_users SET password=?, session_token="" WHERE username=?',
                 (new_hashed, username))
    conn.commit()
    conn.close()
    return jsonify({'ok': True, 'message': '密码重置成功'})

@app.route('/api/user/security/set', methods=['POST'])
def set_security_question():
    """设置安全问题（登录后调用）"""
    data = request.json or {}
    username = data.get('username', '')
    question = data.get('question', '').strip()
    answer = data.get('answer', '').strip()
    if not username or not question or not answer:
        return jsonify({'error': '请填写完整'}), 400
    conn = get_db()
    conn.execute('UPDATE server_users SET security_question=?, security_answer=? WHERE username=?',
                 (question, answer.lower(), username))
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

# ========== 紧急联系人 ==========
@app.route('/api/emergency/contacts', methods=['GET'])
def get_emergency_contacts():
    username = request.args.get('username', '')
    if not username:
        return jsonify({'error': '缺少用户名'}), 400
    conn = get_db()
    rows = conn.execute('SELECT id, name, phone, relation FROM emergency_contacts WHERE username=?', (username,)).fetchall()
    conn.close()
    contacts = [{'id': r['id'], 'name': r['name'], 'phone': r['phone'], 'relation': r['relation']} for r in rows]
    return jsonify({'contacts': contacts})

@app.route('/api/emergency/contacts', methods=['POST'])
def add_emergency_contact():
    data = request.json or {}
    username = data.get('username', '')
    name = data.get('name', '').strip()
    phone = data.get('phone', '').strip()
    relation = data.get('relation', '').strip()
    if not username or not name or not phone:
        return jsonify({'error': '姓名和电话不能为空'}), 400
    conn = get_db()
    conn.execute('INSERT INTO emergency_contacts (username, name, phone, relation) VALUES (?,?,?,?)',
                 (username, name, phone, relation))
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

@app.route('/api/emergency/contacts/delete', methods=['POST'])
def delete_emergency_contact():
    data = request.json or {}
    contact_id = data.get('id')
    username = data.get('username', '')
    if not contact_id or not username:
        return jsonify({'error': '参数不完整'}), 400
    conn = get_db()
    conn.execute('DELETE FROM emergency_contacts WHERE id=? AND username=?', (contact_id, username))
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

@app.route('/api/user/update', methods=['POST'])
def update_user():
    data = request.json or {}
    username = data.get('username', '')
    nickname = data.get('nickname', '')
    avatar = data.get('avatar', '')
    age = data.get('age', None)
    role = data.get('role', '')
    region = data.get('region', '')
    bio = data.get('bio', '')
    conn = get_db()
    if nickname:
        conn.execute('UPDATE server_users SET nickname=? WHERE username=?', (nickname, username))
    if avatar:
        conn.execute('UPDATE server_users SET avatar=? WHERE username=?', (avatar, username))
    if age is not None:
        conn.execute('UPDATE server_users SET age=? WHERE username=?', (age, username))
    if role:
        conn.execute('UPDATE server_users SET role=? WHERE username=?', (role, username))
    if region:
        conn.execute('UPDATE server_users SET region=? WHERE username=?', (region, username))
    if bio:
        conn.execute('UPDATE server_users SET bio=? WHERE username=?', (bio, username))
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

@app.route('/api/admin/ban', methods=['POST'])
@admin_required
def ban_user():
    data = request.json or {}
    username = data.get('username', '')
    conn = get_db()
    conn.execute('UPDATE server_users SET banned=1 WHERE username=?', (username,))
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

@app.route('/api/admin/unban', methods=['POST'])
@admin_required
def unban_user():
    data = request.json or {}
    username = data.get('username', '')
    conn = get_db()
    conn.execute('UPDATE server_users SET banned=0 WHERE username=?', (username,))
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

@app.route('/api/admin/delete_user', methods=['POST'])
@admin_required
def delete_user():
    data = request.json or {}
    username = data.get('username', '')
    if not username:
        return jsonify({'error': '用户名不能为空'}), 400
    conn = get_db()
    conn.execute('DELETE FROM server_users WHERE username=?', (username,))
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

@app.route('/api/admin/delete_device', methods=['POST'])
@admin_required
def delete_device():
    data = request.json or {}
    device_id = data.get('device_id', '')
    if not device_id:
        return jsonify({'error': '设备ID不能为空'}), 400
    conn = get_db()
    conn.execute('DELETE FROM anonymous_posts WHERE device_id=?', (device_id,))
    conn.execute('DELETE FROM anonymous_replies WHERE device_id=?', (device_id,))
    conn.execute('DELETE FROM square_notifications WHERE target_device_id=?', (device_id,))
    conn.execute('DELETE FROM user_events WHERE device_id=?', (device_id,))
    conn.execute('DELETE FROM user_devices WHERE device_id=?', (device_id,))
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

@app.route('/api/admin/delete_post', methods=['POST'])
@admin_required
def admin_delete_post():
    data = request.json or {}
    post_id = data.get('post_id', 0)
    if not post_id:
        return jsonify({'error': '帖子ID不能为空'}), 400
    conn = get_db()
    conn.execute('DELETE FROM anonymous_replies WHERE post_id=?', (post_id,))
    conn.execute('DELETE FROM anonymous_posts WHERE id=?', (post_id,))
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

@app.route('/api/admin/delete_reply', methods=['POST'])
@admin_required
def admin_delete_reply():
    data = request.json or {}
    reply_id = data.get('reply_id', 0)
    if not reply_id:
        return jsonify({'error': '回复ID不能为空'}), 400
    conn = get_db()
    conn.execute('DELETE FROM anonymous_replies WHERE id=?', (reply_id,))
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

@app.route('/api/admin/delete_chat', methods=['POST'])
@admin_required
def admin_delete_chat():
    data = request.json or {}
    msg_id = data.get('msg_id', 0)
    if not msg_id:
        return jsonify({'error': '记录ID不能为空'}), 400
    conn = get_db()
    conn.execute('DELETE FROM chat_history WHERE id=?', (msg_id,))
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

# ========== 管理后台 - 数据管理 ==========

@app.route('/api/admin/records/<table_name>')
@admin_required
def admin_records(table_name):
    valid_tables = ['moods','checkins','chat_history','test_results_sync','anonymous_posts','anonymous_replies','support_messages','emergency_contacts','server_users']
    if table_name not in valid_tables:
        return jsonify({'error': '无效的表名'}), 400
    conn = get_db()
    limit = min(int(request.args.get('limit', 100)), 500)
    offset = int(request.args.get('offset', 0))
    rows = [dict(r) for r in conn.execute(f'SELECT * FROM {table_name} ORDER BY rowid DESC LIMIT ? OFFSET ?', (limit, offset)).fetchall()]
    total = conn.execute(f'SELECT COUNT(*) FROM {table_name}').fetchone()[0]
    conn.close()
    return jsonify({'records': rows, 'total': total})

@app.route('/api/admin/update_record', methods=['POST'])
@admin_required
def admin_update_record():
    data = request.json or {}
    table = data.get('table', '')
    record_id = data.get('id', 0)
    fields = data.get('fields', {})
    valid_tables = {'moods': ['mood','note','gender'], 'checkins': ['day','completed','gender'], 'chat_history': ['role','message','gender'], 'test_results_sync': ['test_title','score','level'], 'anonymous_posts': ['content','anonymous_name'], 'anonymous_replies': ['content','anonymous_name'], 'support_messages': ['content'], 'emergency_contacts': ['name','phone','relation'], 'server_users': ['nickname','gender','banned']}
    if table not in valid_tables or not record_id:
        return jsonify({'error': '参数错误'}), 400
    allowed = valid_tables[table]
    sets = []
    vals = []
    for k, v in fields.items():
        if k in allowed:
            sets.append(f'{k}=?')
            vals.append(v)
    if not sets:
        return jsonify({'error': '无有效字段'}), 400
    vals.append(record_id)
    conn = get_db()
    conn.execute(f'UPDATE {table} SET {",".join(sets)} WHERE id=?', vals)
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

@app.route('/api/admin/delete_record', methods=['POST'])
@admin_required
def admin_delete_record():
    data = request.json or {}
    table = data.get('table', '')
    record_id = data.get('id', 0)
    valid_tables = ['moods','checkins','chat_history','test_results_sync','anonymous_posts','anonymous_replies','support_messages','emergency_contacts']
    if table not in valid_tables or not record_id:
        return jsonify({'error': '参数错误'}), 400
    conn = get_db()
    conn.execute(f'DELETE FROM {table} WHERE id=?', (record_id,))
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

@app.route('/api/admin/delete_support_msg', methods=['POST'])
@admin_required
def admin_delete_support_msg():
    data = request.json or {}
    msg_id = data.get('msg_id', 0)
    if not msg_id:
        return jsonify({'error': '消息ID不能为空'}), 400
    conn = get_db()
    conn.execute('DELETE FROM support_messages WHERE id=?', (msg_id,))
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

@app.route('/api/admin/approve_user', methods=['POST'])
@admin_required
def admin_approve_user():
    data = request.json or {}
    username = data.get('username', '')
    if not username:
        return jsonify({'error': '用户名不能为空'}), 400
    conn = get_db()
    conn.execute('UPDATE server_users SET verify_status="approved" WHERE username=?', (username,))
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

@app.route('/api/admin/reject_user', methods=['POST'])
@admin_required
def admin_reject_user():
    data = request.json or {}
    username = data.get('username', '')
    if not username:
        return jsonify({'error': '用户名不能为空'}), 400
    conn = get_db()
    conn.execute('UPDATE server_users SET verify_status="rejected" WHERE username=?', (username,))
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

# ========== 心情记录接口 ==========
@app.route('/api/moods', methods=['GET'])
def get_moods():
    device_id = request.args.get('device_id', '')
    gender = request.args.get('gender', 'male')
    conn = get_db()
    if device_id:
        moods = conn.execute(
            'SELECT * FROM moods WHERE device_id = ? ORDER BY created_at DESC LIMIT 30',
            (device_id,)
        ).fetchall()
    else:
        moods = conn.execute(
            'SELECT * FROM moods WHERE gender = ? ORDER BY created_at DESC LIMIT 30',
            (gender,)
        ).fetchall()
    conn.close()
    return jsonify({'moods': [dict(m) for m in moods]})

@app.route('/api/moods', methods=['POST'])
def add_mood():
    data = request.get_json()
    conn = get_db()
    conn.execute(
        'INSERT INTO moods (mood, note, gender, device_id) VALUES (?, ?, ?, ?)',
        (data['mood'], data.get('note', ''), data.get('gender', 'male'), data.get('device_id', ''))
    )
    conn.commit()
    conn.close()
    return jsonify({"message": "记录成功"}), 201

# ========== 打卡接口 ==========
@app.route('/api/checkins', methods=['GET'])
def get_checkins():
    gender = request.args.get('gender', 'male')
    conn = get_db()
    checkins = conn.execute(
        'SELECT * FROM checkins WHERE gender = ? ORDER BY created_at DESC LIMIT 30',
        (gender,)
    ).fetchall()
    conn.close()
    return jsonify([dict(c) for c in checkins])

@app.route('/api/checkins', methods=['POST'])
def add_checkin():
    data = request.get_json()
    today = datetime.now().strftime('%Y-%m-%d')
    conn = get_db()
    existing = conn.execute(
        'SELECT * FROM checkins WHERE day = ? AND gender = ?',
        (today, data.get('gender', 'male'))
    ).fetchone()
    if existing:
        conn.close()
        return jsonify({"message": "今天已打卡"}), 200
    conn.execute(
        'INSERT INTO checkins (day, completed, gender) VALUES (?, 1, ?)',
        (today, data.get('gender', 'male'))
    )
    conn.commit()
    conn.close()
    return jsonify({"message": "打卡成功"}), 201

# ========== 美食推荐接口 ==========
@app.route('/api/food', methods=['POST'])
def recommend_food():
    data = request.get_json()
    taste = data.get('taste', '甜')

    food_map = {
        "甜": [
            {"name": "珍珠奶茶", "desc": "甜蜜的奶茶让心情变好", "platforms": ["美团", "饿了么", "抖音团购"]},
            {"name": "芝士蛋糕", "desc": "绵密的芝士口感治愈一切", "platforms": ["美团", "饿了么"]},
            {"name": "草莓大福", "desc": "软糯外皮包裹新鲜草莓", "platforms": ["美团", "饿了么"]},
            {"name": "提拉米苏", "desc": "经典意式甜品，微苦回甘", "platforms": ["美团", "饿了么"]},
        ],
        "辣": [
            {"name": "麻辣烫", "desc": "热辣的汤底驱散所有阴霾", "platforms": ["美团", "饿了么", "抖音团购"]},
            {"name": "重庆小面", "desc": "一碗热辣小面，酣畅淋漓", "platforms": ["美团", "饿了么"]},
            {"name": "螺蛳粉", "desc": "又臭又辣，越吃越上头", "platforms": ["美团", "饿了么"]},
            {"name": "火锅", "desc": "没有什么是一顿火锅解决不了的", "platforms": ["美团", "饿了么", "抖音团购"]},
        ],
        "清淡": [
            {"name": "白粥配小菜", "desc": "清淡暖胃，慢慢治愈", "platforms": ["美团", "饿了么"]},
            {"name": "清汤面", "desc": "简单的面条，温暖的味道", "platforms": ["美团", "饿了么"]},
            {"name": "蒸蛋", "desc": "滑嫩口感，温柔如你", "platforms": ["美团", "饿了么"]},
            {"name": "蔬菜沙拉", "desc": "清爽健康，重新开始", "platforms": ["美团", "饿了么"]},
        ],
        "香": [
            {"name": "烤肉饭", "desc": "香气四溢，满足感爆棚", "platforms": ["美团", "饿了么"]},
            {"name": "炸鸡", "desc": "外酥里嫩，快乐加倍", "platforms": ["美团", "饿了么", "抖音团购"]},
            {"name": "汉堡", "desc": "大口吃肉，烦恼全消", "platforms": ["美团", "饿了么"]},
            {"name": "烤串", "desc": "滋滋冒油，人间值得", "platforms": ["美团", "饿了么"]},
        ],
        "冰": [
            {"name": "冰淇淋", "desc": "冰冰凉凉，心情透心凉", "platforms": ["美团", "饿了么"]},
            {"name": "冰沙", "desc": "夏日必备，清爽解暑", "platforms": ["美团", "饿了么"]},
            {"name": "冰咖啡", "desc": "提神醒脑，重新振作", "platforms": ["美团", "饿了么"]},
            {"name": "冰粉", "desc": "清凉爽滑，甜甜蜜蜜", "platforms": ["美团", "饿了么"]},
        ],
    }

    foods = food_map.get(taste, food_map["甜"])
    return jsonify({"taste": taste, "foods": foods})

# ========== 外卖平台链接 ==========
@app.route('/api/platforms')
def get_platforms():
    platforms = [
        {"name": "美团外卖", "url": "https://waimai.meituan.com", "icon": "Meituan"},
        {"name": "饿了么", "url": "https://www.ele.me", "icon": "Eleme"},
        {"name": "抖音团购", "url": "https://www.douyin.com", "icon": "Douyin"},
        {"name": "大众点评", "url": "https://www.dianping.com", "icon": "Dianping"},
    ]
    return jsonify(platforms)

# ========== AI情感助手（调用小米mimo API） ==========
@app.route('/api/chat', methods=['POST'])
@limiter.limit("10 per minute")
def chat():
    data = request.get_json()
    user_message = data.get('message', '')
    gender = data.get('gender', 'male')

    # 保存用户消息
    conn = get_db()
    conn.execute(
        'INSERT INTO chat_history (role, message, gender) VALUES (?, ?, ?)',
        ('user', user_message, gender)
    )
    conn.commit()

    # 获取历史对话
    history = conn.execute(
        'SELECT role, message FROM chat_history WHERE gender = ? ORDER BY created_at DESC LIMIT 10',
        (gender,)
    ).fetchall()
    conn.close()

    # 构建消息列表
    system_prompt = """你是一个温暖贴心的AI情感助手，名叫"心语"。
你的任务是倾听用户的心声，给予温暖的回应和实用的建议。
你会：
1. 认真倾听，表示理解和共情
2. 给出具体可行的缓解情绪的小方法
3. 语气温柔亲切，像朋友一样
4. 如果用户情绪很严重，温柔地建议寻求专业帮助
5. 回复简洁温暖，不要太长"""

    messages = []
    for h in reversed(history):
        role = h['role']
        if role not in ('user', 'assistant'):
            role = 'user'
        messages.append({"role": role, "content": h['message']})

    # 调用小米mimo API
    try:
        api_url = "https://token-plan-cn.xiaomimimo.com/anthropic/v1/messages"
        headers = {
            "Content-Type": "application/json",
            "x-api-key": os.environ.get('MIMO_API_KEY', ''),
            "anthropic-version": "2023-06-01"
        }
        payload = {
            "model": "mimo-v2.5",
            "system": system_prompt,
            "max_tokens": 1024,
            "messages": messages
        }
        resp = requests.post(api_url, headers=headers, json=payload, timeout=30)
        result = resp.json()
        content_list = result.get('content', [])
        ai_message = ''
        for item in content_list:
            if item.get('type') == 'text':
                ai_message = item.get('text', '')
                break
            elif item.get('type') == 'thinking':
                ai_message = item.get('thinking', '')
        if not ai_message:
            ai_message = '抱歉，我现在有点累了，稍后再聊好吗？'
    except Exception as e:
        ai_message = "抱歉，AI助手暂时不可用。但请记住，你并不孤单，随时可以拨打心理援助热线：12320-5"

    # 保存AI回复
    conn = get_db()
    conn.execute(
        'INSERT INTO chat_history (role, message, gender) VALUES (?, ?, ?)',
        ('assistant', ai_message, gender)
    )
    conn.commit()
    conn.close()

    return jsonify({"reply": ai_message})

@app.route('/api/chat/history')
def chat_history():
    gender = request.args.get('gender', 'male')
    conn = get_db()
    history = conn.execute(
        'SELECT role, message, created_at FROM chat_history WHERE gender = ? ORDER BY created_at DESC LIMIT 50',
        (gender,)
    ).fetchall()
    conn.close()
    return jsonify([dict(h) for h in reversed(history)])

# ========== 天气代理 ==========
def _get_city_from_coords(lat, lon):
    """通过经纬度获取城市名，使用多个服务容错"""
    # 方法1：photon.komoot.io（免费，对中国城市准确）
    try:
        r = requests.get(
            f'https://photon.komoot.io/reverse?lat={lat}&lon={lon}&limit=1',
            headers={'User-Agent': 'Mozilla/5.0'},
            timeout=8
        ).json()
        features = r.get('features', [])
        if features:
            props = features[0].get('properties', {})
            city = props.get('city') or props.get('town') or props.get('village') or ''
            state = props.get('state', '')
            # 中国城市：city 字段可能是区/县名，需要取 state（直辖市/省名包含城市名）
            if city and any(s in city for s in ['区', '县', '旗']):
                city = state or city
            if not city and state:
                city = state
            if city:
                return city
    except Exception:
        pass
    # 方法2：Nominatim 兜底
    try:
        r = requests.get(
            f'https://nominatim.openstreetmap.org/reverse?lat={lat}&lon={lon}&format=json&accept-language=zh',
            headers={'User-Agent': 'XinyuApp/1.0'},
            timeout=8
        ).json()
        addr = r.get('address', {})
        city = addr.get('city') or addr.get('town') or addr.get('village') or addr.get('county') or ''
        if city and any(s in city for s in ['区', '县', '旗']):
            city = addr.get('state', '') or city
        if city:
            return city
    except Exception:
        pass
    return '北京'

@app.route('/api/weather')
def proxy_weather():
    city = request.args.get('city', '')
    lat = request.args.get('lat', type=float)
    lon = request.args.get('lon', type=float)
    try:
        import urllib.parse
        display_name = city

        if lat is not None and lon is not None:
            # 直接用经纬度查天气（最准确）
            display_name = _get_city_from_coords(lat, lon)
        elif city:
            geo_url = 'https://geocoding-api.open-meteo.com/v1/search?name=' + urllib.parse.quote(city) + '&count=5&language=zh'
            geo_data = requests.get(geo_url, timeout=10).json()
            results = geo_data.get('results', [])
            if not results:
                return jsonify({'error': '找不到该城市'}), 404
            # 优先选中国的结果
            loc = None
            for r in results:
                if r.get('country_code') in ('CN', 'cn'):
                    loc = r
                    break
            if not loc:
                loc = results[0]
            lat, lon = loc['latitude'], loc['longitude']
            display_name = loc.get('name', city)
        else:
            lat, lon = 39.9042, 116.4074
            display_name = '北京'

        weather_url = (f'https://api.open-meteo.com/v1/forecast?latitude={lat}&longitude={lon}'
                       f'&current=temperature_2m,relative_humidity_2m,weather_code,wind_speed_10m&language=zh')
        w = requests.get(weather_url, timeout=10).json()['current']

        WMO_DESC = {0:'晴',1:'大部晴朗',2:'多云',3:'阴天',45:'雾',48:'雾',51:'毛毛雨',53:'毛毛雨',55:'毛毛雨',
                     61:'小雨',63:'中雨',65:'大雨',71:'小雪',73:'中雪',75:'大雪',80:'阵雨',81:'中阵雨',82:'大阵雨',
                     95:'雷暴',96:'雷暴冰雹',99:'雷暴冰雹'}
        WMO_ICON = {0:'☀️',1:'🌤️',2:'⛅',3:'☁️',45:'🌫️',48:'🌫️',51:'🌧️',53:'🌧️',55:'🌧️',
                     61:'🌧️',63:'🌧️',65:'🌧️',71:'❄️',73:'❄️',75:'❄️',80:'🌦️',81:'🌦️',82:'🌦️',
                     95:'⛈️',96:'⛈️',99:'⛈️'}
        code = w['weather_code']
        return jsonify({
            'city': display_name,
            'temp': str(int(w['temperature_2m'])) + '°C',
            'desc': WMO_DESC.get(code, '未知'),
            'icon': WMO_ICON.get(code, '🌤️'),
            'detail': f"湿度 {w['relative_humidity_2m']}% · 风速 {w['wind_speed_10m']}km/h"
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# ========== 匿名广场 ==========
ANIMALS = ['小猫🐱','小狗🐶','小兔🐰','小熊🐻','小狐🦊','小熊猫🐼','小牛🐮','小猪🐷','小猴🐵','小狮🦁','小虎🐯','小蛙🐸','小鱼🐟','小鸟🐦','小鹿🦌','小鸭🦆']
ADJS = ['快乐的','温柔的','勇敢的','安静的','活泼的','聪明的','可爱的','治愈的','善良的','元气的','柔软的','温暖的']

def get_anonymous_name(device_id):
    h = hash(device_id)
    animal = ANIMALS[h % len(ANIMALS)]
    adj = ADJS[(h // len(ANIMALS)) % len(ADJS)]
    return adj + animal

def get_real_name(device_id):
    """获取真实用户名，优先从user_devices查找"""
    conn = get_db()
    row = conn.execute('SELECT nickname FROM user_devices WHERE device_id=?', (device_id,)).fetchone()
    conn.close()
    if row and row['nickname']:
        return row['nickname']
    return get_anonymous_name(device_id)

def init_square_tables():
    conn = get_db()
    conn.execute('''CREATE TABLE IF NOT EXISTS anonymous_posts (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        content TEXT NOT NULL,
        anonymous_name TEXT,
        device_id TEXT,
        media_url TEXT DEFAULT '',
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')
    # Add media_url column if missing (for existing DBs)
    try:
        conn.execute('ALTER TABLE anonymous_posts ADD COLUMN media_url TEXT DEFAULT ""')
    except:
        pass
    # Add tags column for story filtering
    try:
        conn.execute('ALTER TABLE anonymous_posts ADD COLUMN tags TEXT DEFAULT "[]"')
    except:
        pass
    conn.execute('''CREATE TABLE IF NOT EXISTS anonymous_replies (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        post_id INTEGER NOT NULL,
        content TEXT NOT NULL,
        anonymous_name TEXT,
        device_id TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')
    conn.execute('''CREATE TABLE IF NOT EXISTS square_notifications (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        target_device_id TEXT NOT NULL,
        from_anonymous_name TEXT,
        post_content TEXT,
        reply_content TEXT,
        is_read INTEGER DEFAULT 0,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')
    conn.execute('''CREATE TABLE IF NOT EXISTS anonymous_likes (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        post_id INTEGER NOT NULL,
        device_id TEXT NOT NULL,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        UNIQUE(post_id, device_id)
    )''')
    # source字段区分筑爱故事汇和叙心树洞
    try:
        conn.execute('ALTER TABLE anonymous_posts ADD COLUMN source TEXT DEFAULT "treehole"')
    except:
        pass
    # 投递树洞案例表
    conn.execute('''CREATE TABLE IF NOT EXISTS treehole_cases (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        content TEXT NOT NULL,
        author_device TEXT,
        author_username TEXT DEFAULT '',
        author_type TEXT DEFAULT 'user',
        tags TEXT DEFAULT '[]',
        status TEXT DEFAULT 'open',
        like_count INTEGER DEFAULT 0,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')
    # 投递树洞辅导师回复表
    conn.execute('''CREATE TABLE IF NOT EXISTS treehole_answers (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        case_id INTEGER NOT NULL,
        content TEXT NOT NULL,
        author_username TEXT DEFAULT '',
        author_type TEXT DEFAULT 'counselor',
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')
    conn.execute('''CREATE TABLE IF NOT EXISTS counselor_activities (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        title TEXT NOT NULL,
        content TEXT NOT NULL,
        activity_type TEXT DEFAULT 'offline',
        location TEXT DEFAULT '',
        activity_time TEXT DEFAULT '',
        author_username TEXT DEFAULT '',
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')
    conn.commit()
    conn.close()

init_square_tables()

# ===== 测评结果同步 =====
def init_assessment_tables():
    conn = get_db()
    conn.execute('''CREATE TABLE IF NOT EXISTS test_results_sync (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT NOT NULL,
        test_id TEXT NOT NULL,
        test_title TEXT NOT NULL,
        score INTEGER NOT NULL,
        level TEXT,
        description TEXT,
        created_at INTEGER,
        synced_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        UNIQUE(username, test_id, created_at)
    )''')
    conn.commit()
    conn.close()

init_assessment_tables()

@app.route('/api/assessment/sync', methods=['POST'])
def sync_assessment():
    data = request.json or {}
    username = data.get('username', '')
    test_id = data.get('test_id', '')
    test_title = data.get('test_title', '')
    score = data.get('score', 0)
    level = data.get('level', '')
    description = data.get('description', '')
    created_at = data.get('created_at', 0)
    if not username or not test_id:
        return jsonify({'error': '参数缺失'}), 400
    conn = get_db()
    conn.execute('''INSERT OR REPLACE INTO test_results_sync
        (username, test_id, test_title, score, level, description, created_at)
        VALUES (?, ?, ?, ?, ?, ?, ?)''',
        (username, test_id, test_title, score, level, description, created_at))
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

@app.route('/api/assessment/results', methods=['GET'])
def get_assessment_results():
    username = request.args.get('username', '')
    if not username:
        return jsonify({'results': []})
    conn = get_db()
    rows = conn.execute(
        'SELECT * FROM test_results_sync WHERE username=? ORDER BY created_at DESC LIMIT 100',
        (username,)).fetchall()
    conn.close()
    return jsonify({'results': [dict(r) for r in rows]})

@app.route('/api/square/posts', methods=['GET'])
def get_square_posts():
    conn = get_db()
    device_id = request.args.get('device_id', '')
    sort = request.args.get('sort', 'time')  # 'time' or 'hot'
    tag = request.args.get('tag', '')
    search = request.args.get('search', '')
    source = request.args.get('source', '')  # 'story' or 'treehole'
    order = 'p.created_at DESC' if sort == 'time' else '(like_count + reply_count * 2) DESC, p.created_at DESC'

    where_clause = ''
    params = []
    if source:
        where_clause = "WHERE p.source=?"
        params.append(source)
    if tag:
        where_clause = ("WHERE " if not where_clause else where_clause + " AND ") + "p.tags LIKE ?"
        params.append(f'%"{tag}"%')
    if search:
        where_clause = ("WHERE " if not where_clause else where_clause + " AND ") + "p.content LIKE ?"
        params.append(f'%{search}%')

    posts = conn.execute(
        f'''SELECT p.id, p.content, p.anonymous_name, p.device_id,
           p.media_url, p.tags,
           strftime('%s', p.created_at) as created_at,
           (SELECT COUNT(*) FROM anonymous_replies r WHERE r.post_id=p.id) as reply_count,
           (SELECT COUNT(*) FROM anonymous_likes l WHERE l.post_id=p.id) as like_count
           FROM anonymous_posts p {where_clause} ORDER BY {order} LIMIT 50''',
        params
    ).fetchall()
    # Check which posts the current device has liked
    liked_posts = set()
    if device_id:
        rows = conn.execute('SELECT post_id FROM anonymous_likes WHERE device_id=?', (device_id,)).fetchall()
        liked_posts = {r['post_id'] for r in rows}
    conn.close()
    result = []
    for p in posts:
        d = dict(p)
        d['liked'] = p['id'] in liked_posts
        try:
            d['tags'] = json.loads(d['tags']) if d['tags'] else []
        except:
            d['tags'] = []
        result.append(d)
    return jsonify({'posts': result})

@app.route('/api/square/posts', methods=['POST'])
@limiter.limit("5 per minute")
def create_square_post():
    data = request.json or {}
    content = data.get('content', '').strip()
    device_id = data.get('device_id', '')
    media_url = data.get('media_url', '')
    tags = json.dumps(data.get('tags', []), ensure_ascii=False)
    if not content and not media_url:
        return jsonify({'error': '内容不能为空'}), 400
    # 筑爱故事汇：只有辅导师和志愿者可以发布帖子
    username = data.get('username', '')
    if username:
        conn = get_db()
        user = conn.execute('SELECT user_type FROM server_users WHERE username=?', (username,)).fetchone()
        conn.close()
        if not user or user['user_type'] == 'user':
            return jsonify({'error': '只有辅导师和志愿者才能发布帖子，请先申请成为志愿者或辅导师'}), 403
    else:
        return jsonify({'error': '请先登录'}), 401
    anon_name = get_real_name(device_id)
    source = data.get('source', 'treehole')
    conn = get_db()
    conn.execute('INSERT INTO anonymous_posts (content, anonymous_name, device_id, media_url, tags, source) VALUES (?,?,?,?,?,?)',
                 (content, anon_name, device_id, media_url, tags, source))
    conn.commit()
    post_id = conn.execute('SELECT last_insert_rowid()').fetchone()[0]
    conn.close()
    return jsonify({'ok': True, 'id': post_id, 'anonymous_name': anon_name})

@app.route('/api/square/posts/<int:post_id>/like', methods=['POST'])
@limiter.limit("30 per minute")
def toggle_like(post_id):
    data = request.json or {}
    device_id = data.get('device_id', '')
    if not device_id:
        return jsonify({'error': '设备ID不能为空'}), 400
    conn = get_db()
    existing = conn.execute('SELECT id FROM anonymous_likes WHERE post_id=? AND device_id=?', (post_id, device_id)).fetchone()
    if existing:
        conn.execute('DELETE FROM anonymous_likes WHERE post_id=? AND device_id=?', (post_id, device_id))
        liked = False
    else:
        conn.execute('INSERT INTO anonymous_likes (post_id, device_id) VALUES (?,?)', (post_id, device_id))
        liked = True
    conn.commit()
    count = conn.execute('SELECT COUNT(*) FROM anonymous_likes WHERE post_id=?', (post_id,)).fetchone()[0]
    conn.close()
    return jsonify({'ok': True, 'liked': liked, 'like_count': count})

@app.route('/api/square/posts/<int:post_id>/replies', methods=['GET'])
def get_square_replies(post_id):
    conn = get_db()
    replies = conn.execute(
        "SELECT id, content, anonymous_name, strftime('%s', created_at) as created_at FROM anonymous_replies WHERE post_id=? ORDER BY created_at ASC",
        (post_id,)
    ).fetchall()
    conn.close()
    return jsonify({'replies': [dict(r) for r in replies]})

@app.route('/api/square/posts/<int:post_id>/replies', methods=['POST'])
@limiter.limit("10 per minute")
def create_square_reply(post_id):
    data = request.json or {}
    content = data.get('content', '').strip()
    device_id = data.get('device_id', '')
    if not content:
        return jsonify({'error': '内容不能为空'}), 400
    anon_name = get_real_name(device_id)
    conn = get_db()
    conn.execute('INSERT INTO anonymous_replies (post_id, content, anonymous_name, device_id) VALUES (?,?,?,?)',
                 (post_id, content, anon_name, device_id))
    # Create notification for post author
    post = conn.execute('SELECT device_id, content FROM anonymous_posts WHERE id=?', (post_id,)).fetchone()
    if post and post['device_id'] and post['device_id'] != device_id:
        conn.execute('INSERT INTO square_notifications (target_device_id, from_anonymous_name, post_content, reply_content) VALUES (?,?,?,?)',
                     (post['device_id'], anon_name, post['content'], content))
    conn.commit()
    conn.close()
    return jsonify({'ok': True, 'anonymous_name': anon_name})

@app.route('/api/square/posts/<int:post_id>', methods=['DELETE'])
def delete_square_post(post_id):
    data = request.json or {}
    device_id = data.get('device_id', '')
    conn = get_db()
    post = conn.execute('SELECT device_id FROM anonymous_posts WHERE id=?', (post_id,)).fetchone()
    if not post:
        conn.close()
        return jsonify({'error': '帖子不存在'}), 404
    if post['device_id'] != device_id:
        conn.close()
        return jsonify({'error': '只能删除自己的帖子'}), 403
    conn.execute('DELETE FROM anonymous_replies WHERE post_id=?', (post_id,))
    conn.execute('DELETE FROM anonymous_posts WHERE id=?', (post_id,))
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

@app.route('/api/square/replies/<int:reply_id>', methods=['DELETE'])
def delete_square_reply(reply_id):
    data = request.json or {}
    device_id = data.get('device_id', '')
    conn = get_db()
    reply = conn.execute('SELECT device_id FROM anonymous_replies WHERE id=?', (reply_id,)).fetchone()
    if not reply:
        conn.close()
        return jsonify({'error': '回复不存在'}), 404
    if reply['device_id'] != device_id:
        conn.close()
        return jsonify({'error': '只能删除自己的回复'}), 403
    conn.execute('DELETE FROM anonymous_replies WHERE id=?', (reply_id,))
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

@app.route('/api/square/notifications', methods=['GET'])
def get_square_notifications():
    device_id = request.args.get('device_id', '')
    if not device_id:
        return jsonify({'notifications': [], 'unread': 0})
    conn = get_db()
    notifs = conn.execute(
        'SELECT id, from_anonymous_name, post_content, reply_content, is_read, created_at FROM square_notifications WHERE target_device_id=? ORDER BY created_at DESC LIMIT 20',
        (device_id,)
    ).fetchall()
    unread = conn.execute(
        'SELECT COUNT(*) FROM square_notifications WHERE target_device_id=? AND is_read=0',
        (device_id,)
    ).fetchone()[0]
    conn.close()
    return jsonify({'notifications': [dict(n) for n in notifs], 'unread': unread})

@app.route('/api/square/notifications/read', methods=['POST'])
def mark_notifications_read():
    data = request.json or {}
    device_id = data.get('device_id', '')
    conn = get_db()
    conn.execute('UPDATE square_notifications SET is_read=1 WHERE target_device_id=?', (device_id,))
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

# ========== 投递树洞 ==========
@app.route('/api/treehole/cases', methods=['GET'])
def get_treehole_cases():
    sort = request.args.get('sort', 'hot')
    conn = get_db()
    if sort == 'hot':
        rows = conn.execute('SELECT * FROM treehole_cases ORDER BY like_count DESC, created_at DESC LIMIT 50').fetchall()
    else:
        rows = conn.execute('SELECT * FROM treehole_cases ORDER BY created_at DESC LIMIT 50').fetchall()
    cases = []
    for r in rows:
        answer_count = conn.execute('SELECT COUNT(*) as c FROM treehole_answers WHERE case_id=?', (r['id'],)).fetchone()['c']
        cases.append({
            'id': r['id'], 'content': r['content'], 'author_device': r['author_device'],
            'author_username': r['author_username'], 'author_type': r['author_type'],
            'tags': r['tags'], 'status': r['status'], 'like_count': r['like_count'],
            'answer_count': answer_count, 'created_at': r['created_at']
        })
    conn.close()
    return jsonify({'cases': cases})

@app.route('/api/treehole/cases', methods=['POST'])
@limiter.limit("5 per minute")
def create_treehole_case():
    data = request.json or {}
    content = data.get('content', '').strip()
    device_id = data.get('device_id', '')
    username = data.get('username', '')
    author_type = data.get('author_type', 'user')
    tags = json.dumps(data.get('tags', []), ensure_ascii=False)
    if not content:
        return jsonify({'error': '请输入内容'}), 400
    conn = get_db()
    conn.execute('INSERT INTO treehole_cases (content, author_device, author_username, author_type, tags) VALUES (?,?,?,?,?)',
                 (content, device_id, username, author_type, tags))
    conn.commit()
    case_id = conn.execute('SELECT last_insert_rowid()').fetchone()[0]
    conn.close()
    return jsonify({'ok': True, 'id': case_id})

@app.route('/api/treehole/cases/<int:case_id>/like', methods=['POST'])
@limiter.limit("30 per minute")
def like_treehole_case(case_id):
    data = request.json or {}
    device_id = data.get('device_id', '')
    conn = get_db()
    existing = conn.execute('SELECT id FROM treehole_case_likes WHERE case_id=? AND device_id=?', (case_id, device_id)).fetchone()
    if existing:
        conn.execute('DELETE FROM treehole_case_likes WHERE case_id=? AND device_id=?', (case_id, device_id))
        conn.execute('UPDATE treehole_cases SET like_count=MAX(0,like_count-1) WHERE id=?', (case_id,))
    else:
        try:
            conn.execute('INSERT INTO treehole_case_likes (case_id, device_id) VALUES (?,?)', (case_id, device_id))
            conn.execute('UPDATE treehole_cases SET like_count=like_count+1 WHERE id=?', (case_id,))
        except:
            pass
    conn.commit()
    row = conn.execute('SELECT like_count FROM treehole_cases WHERE id=?', (case_id,)).fetchone()
    conn.close()
    return jsonify({'ok': True, 'like_count': row['like_count'] if row else 0})

@app.route('/api/treehole/cases/<int:case_id>/answers', methods=['GET'])
def get_treehole_answers(case_id):
    conn = get_db()
    rows = conn.execute('SELECT * FROM treehole_answers WHERE case_id=? ORDER BY created_at ASC', (case_id,)).fetchall()
    answers = [{'id': r['id'], 'content': r['content'], 'author_username': r['author_username'],
                'author_type': r['author_type'], 'created_at': r['created_at']} for r in rows]
    conn.close()
    return jsonify({'answers': answers})

@app.route('/api/treehole/cases/<int:case_id>/answers', methods=['POST'])
@limiter.limit("10 per minute")
def create_treehole_answer(case_id):
    data = request.json or {}
    content = data.get('content', '').strip()
    username = data.get('username', '')
    author_type = data.get('author_type', 'counselor')
    if not content:
        return jsonify({'error': '请输入回复内容'}), 400
    conn = get_db()
    # 更新案例状态为已回复
    conn.execute('UPDATE treehole_cases SET status="answered" WHERE id=?', (case_id,))
    conn.execute('INSERT INTO treehole_answers (case_id, content, author_username, author_type) VALUES (?,?,?,?)',
                 (case_id, content, username, author_type))
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

# 创建案例点赞表
def _init_treehole_likes():
    conn = get_db()
    conn.execute('''CREATE TABLE IF NOT EXISTS treehole_case_likes (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        case_id INTEGER NOT NULL,
        device_id TEXT NOT NULL,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        UNIQUE(case_id, device_id)
    )''')
    conn.commit()
    conn.close()
_init_treehole_likes()

# ========== 辅导师活动 ==========
@app.route('/api/activities', methods=['GET'])
def get_activities():
    conn = get_db()
    rows = conn.execute('SELECT * FROM counselor_activities ORDER BY created_at DESC LIMIT 50').fetchall()
    conn.close()
    return jsonify({'activities': [dict(r) for r in rows]})

@app.route('/api/activities', methods=['POST'])
@limiter.limit("10 per minute")
def create_activity():
    data = request.json or {}
    title = data.get('title', '').strip()
    content = data.get('content', '').strip()
    activity_type = data.get('activity_type', 'offline')
    location = data.get('location', '').strip()
    activity_time = data.get('activity_time', '').strip()
    username = data.get('username', '')
    if not title or not content:
        return jsonify({'error': '请填写活动标题和内容'}), 400
    # 检查是否为辅导师
    conn = get_db()
    user = conn.execute('SELECT user_type FROM server_users WHERE username=?', (username,)).fetchone()
    if not user or user['user_type'] != 'counselor':
        conn.close()
        return jsonify({'error': '只有心理辅导师才能创建活动'}), 403
    conn.execute('INSERT INTO counselor_activities (title, content, activity_type, location, activity_time, author_username) VALUES (?,?,?,?,?,?)',
                 (title, content, activity_type, location, activity_time, username))
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

@app.route('/api/activities/<int:act_id>', methods=['DELETE'])
def delete_activity(act_id):
    data = request.json or {}
    username = data.get('username', '')
    conn = get_db()
    act = conn.execute('SELECT author_username FROM counselor_activities WHERE id=?', (act_id,)).fetchone()
    if not act:
        conn.close()
        return jsonify({'error': '活动不存在'}), 404
    # 仅创建者或管理员可删除
    user = conn.execute('SELECT user_type FROM server_users WHERE username=?', (username,)).fetchone()
    if act['author_username'] != username and (not user or user['user_type'] != 'admin'):
        conn.close()
        return jsonify({'error': '无权删除'}), 403
    conn.execute('DELETE FROM counselor_activities WHERE id=?', (act_id,))
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

# ========== 数据统计面板 ==========
@app.route('/api/stats')
def get_stats():
    gender = request.args.get('gender', 'male')
    conn = get_db()

    # 心情记录总数
    mood_count = conn.execute(
        'SELECT COUNT(*) as c FROM moods WHERE gender = ?', (gender,)
    ).fetchone()['c']

    # 打卡天数
    checkin_count = conn.execute(
        'SELECT COUNT(*) as c FROM checkins WHERE gender = ? AND completed = 1', (gender,)
    ).fetchone()['c']

    # 各心情统计
    mood_stats = conn.execute(
        'SELECT mood, COUNT(*) as count FROM moods WHERE gender = ? GROUP BY mood ORDER BY count DESC',
        (gender,)
    ).fetchall()

    # 最近7天打卡情况
    recent_checkins = conn.execute(
        "SELECT day, completed FROM checkins WHERE gender = ? AND day >= date('now', '-7 days') ORDER BY day",
        (gender,)
    ).fetchall()

    conn.close()

    return jsonify({
        "mood_count": mood_count,
        "checkin_count": checkin_count,
        "mood_stats": [dict(m) for m in mood_stats],
        "recent_checkins": [dict(c) for c in recent_checkins]
    })

# ========== 心理援助电话 ==========
@app.route('/api/helplines')
def get_helplines():
    helplines = [
        {"name": "全国心理援助热线", "phone": "12320-5", "desc": "24小时免费心理危机干预"},
        {"name": "北京心理危机研究与干预中心", "phone": "010-82951332", "desc": "24小时心理危机干预"},
        {"name": "希望24热线", "phone": "400-161-9995", "desc": "24小时生命教育与危机干预"},
        {"name": "青少年心理援助热线", "phone": "12355", "desc": "共青团青少年服务热线"},
        {"name": "女性心理援助热线", "phone": "12338", "desc": "全国妇联妇女维权热线"},
        {"name": "生命热线", "phone": "400-821-1215", "desc": "上海市精神卫生中心"},
    ]
    return jsonify(helplines)

# ========== 学生管理系统接口 ==========

# 获取所有学生
@app.route('/api/students', methods=['GET'])
def get_students():
    search = request.args.get('search', '')
    conn = get_db()
    if search:
        students = conn.execute(
            'SELECT * FROM students WHERE name LIKE ? OR student_id LIKE ? OR major LIKE ? ORDER BY created_at DESC',
            (f'%{search}%', f'%{search}%', f'%{search}%')
        ).fetchall()
    else:
        students = conn.execute('SELECT * FROM students ORDER BY created_at DESC').fetchall()
    conn.close()
    return jsonify([dict(s) for s in students])

# 获取单个学生
@app.route('/api/students/<int:id>', methods=['GET'])
def get_student(id):
    conn = get_db()
    student = conn.execute('SELECT * FROM students WHERE id = ?', (id,)).fetchone()
    conn.close()
    if student:
        return jsonify(dict(student))
    return jsonify({"error": "学生不存在"}), 404

# 添加学生
@app.route('/api/students', methods=['POST'])
def add_student():
    data = request.get_json()
    conn = get_db()
    try:
        conn.execute(
            '''INSERT INTO students (student_id, name, gender, age, major, class_name, phone, email, address)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)''',
            (data['student_id'], data['name'], data.get('gender', ''),
             data.get('age'), data.get('major', ''), data.get('class_name', ''),
             data.get('phone', ''), data.get('email', ''), data.get('address', ''))
        )
        conn.commit()
        new_id = conn.execute('SELECT last_insert_rowid()').fetchone()[0]
        conn.close()
        return jsonify({"id": new_id, "message": "添加成功"}), 201
    except sqlite3.IntegrityError:
        conn.close()
        return jsonify({"error": "学号已存在"}), 400

# 修改学生
@app.route('/api/students/<int:id>', methods=['PUT'])
def update_student(id):
    data = request.get_json()
    conn = get_db()
    student = conn.execute('SELECT * FROM students WHERE id = ?', (id,)).fetchone()
    if not student:
        conn.close()
        return jsonify({"error": "学生不存在"}), 404

    conn.execute(
        '''UPDATE students SET student_id=?, name=?, gender=?, age=?, major=?, class_name=?, phone=?, email=?, address=?
           WHERE id=?''',
        (data.get('student_id', student['student_id']),
         data.get('name', student['name']),
         data.get('gender', student['gender']),
         data.get('age', student['age']),
         data.get('major', student['major']),
         data.get('class_name', student['class_name']),
         data.get('phone', student['phone']),
         data.get('email', student['email']),
         data.get('address', student['address']),
         id)
    )
    conn.commit()
    conn.close()
    return jsonify({"message": "修改成功"})

# 删除学生
@app.route('/api/students/<int:id>', methods=['DELETE'])
def delete_student(id):
    conn = get_db()
    student = conn.execute('SELECT * FROM students WHERE id = ?', (id,)).fetchone()
    if not student:
        conn.close()
        return jsonify({"error": "学生不存在"}), 404
    conn.execute('DELETE FROM students WHERE id = ?', (id,))
    conn.commit()
    conn.close()
    return jsonify({"message": "删除成功"})

# 学生统计
@app.route('/api/students/stats')
def student_stats():
    conn = get_db()
    total = conn.execute('SELECT COUNT(*) as c FROM students').fetchone()['c']
    male = conn.execute("SELECT COUNT(*) as c FROM students WHERE gender='男'").fetchone()['c']
    female = conn.execute("SELECT COUNT(*) as c FROM students WHERE gender='女'").fetchone()['c']
    majors = conn.execute('SELECT major, COUNT(*) as count FROM students WHERE major != "" GROUP BY major ORDER BY count DESC').fetchall()
    conn.close()
    return jsonify({
        "total": total,
        "male": male,
        "female": female,
        "majors": [dict(m) for m in majors]
    })

# ========== 数据库管理接口 ==========

# 获取所有表信息
@app.route('/api/db/tables')
def get_tables():
    conn = get_db()
    tables = conn.execute("SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%'").fetchall()
    result = []
    for table in tables:
        name = table['name']
        count = conn.execute(f'SELECT COUNT(*) as c FROM {name}').fetchone()['c']
        columns = conn.execute(f'PRAGMA table_info({name})').fetchall()
        result.append({
            'name': name,
            'count': count,
            'columns': [{'name': c['name'], 'type': c['type']} for c in columns]
        })
    conn.close()
    return jsonify(result)

# 获取表数据
@app.route('/api/db/table/<table_name>')
def get_table_data(table_name):
    # 安全校验
    conn = get_db()
    valid_tables = [t['name'] for t in conn.execute("SELECT name FROM sqlite_master WHERE type='table'").fetchall()]
    if table_name not in valid_tables:
        conn.close()
        return jsonify({"error": "无效的表名"}), 400

    page = int(request.args.get('page', 1))
    per_page = int(request.args.get('per_page', 20))
    offset = (page - 1) * per_page

    total = conn.execute(f'SELECT COUNT(*) as c FROM {table_name}').fetchone()['c']
    data = conn.execute(f'SELECT * FROM {table_name} ORDER BY rowid DESC LIMIT ? OFFSET ?', (per_page, offset)).fetchall()
    columns = conn.execute(f'PRAGMA table_info({table_name})').fetchall()
    conn.close()

    return jsonify({
        'table': table_name,
        'columns': [c['name'] for c in columns],
        'data': [dict(row) for row in data],
        'total': total,
        'page': page,
        'per_page': per_page,
        'total_pages': (total + per_page - 1) // per_page
    })

# 删除记录
@app.route('/api/db/table/<table_name>/<int:row_id>', methods=['DELETE'])
def delete_row(table_name, row_id):
    conn = get_db()
    valid_tables = [t['name'] for t in conn.execute("SELECT name FROM sqlite_master WHERE type='table'").fetchall()]
    if table_name not in valid_tables:
        conn.close()
        return jsonify({"error": "无效的表名"}), 400

    conn.execute(f'DELETE FROM {table_name} WHERE id = ?', (row_id,))
    conn.commit()
    conn.close()
    return jsonify({"message": "删除成功"})

# 清空表
@app.route('/api/db/table/<table_name>', methods=['DELETE'])
def clear_table(table_name):
    conn = get_db()
    valid_tables = [t['name'] for t in conn.execute("SELECT name FROM sqlite_master WHERE type='table'").fetchall()]
    if table_name not in valid_tables:
        conn.close()
        return jsonify({"error": "无效的表名"}), 400

    conn.execute(f'DELETE FROM {table_name}')
    conn.commit()
    conn.close()
    return jsonify({"message": f"已清空 {table_name} 表"})

# 执行自定义SQL查询（只允许SELECT）
@app.route('/api/db/query', methods=['POST'])
def execute_query():
    data = request.get_json()
    sql = data.get('sql', '').strip()

    # 安全校验：只允许SELECT
    if not sql.upper().startswith('SELECT'):
        return jsonify({"error": "只允许执行SELECT查询"}), 400

    # 禁止危险操作
    dangerous = ['DROP', 'DELETE', 'UPDATE', 'INSERT', 'ALTER', 'CREATE', 'TRUNCATE']
    for word in dangerous:
        if word in sql.upper():
            return jsonify({"error": f"禁止执行{word}操作"}), 400

    try:
        conn = get_db()
        result = conn.execute(sql).fetchall()
        columns = conn.execute(sql).description
        conn.close()

        return jsonify({
            'columns': [c[0] for c in columns] if columns else [],
            'data': [dict(row) for row in result],
            'count': len(result)
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 400

# 数据库统计信息
@app.route('/api/db/stats')
def db_stats():
    conn = get_db()
    tables = conn.execute("SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%'").fetchall()

    stats = []
    total_records = 0
    for table in tables:
        name = table['name']
        count = conn.execute(f'SELECT COUNT(*) as c FROM {name}').fetchone()['c']
        total_records += count
        stats.append({'table': name, 'count': count})

    # 数据库文件大小
    db_size = os.path.getsize(DB_PATH) if os.path.exists(DB_PATH) else 0
    conn.close()

    return jsonify({
        'tables': stats,
        'total_tables': len(stats),
        'total_records': total_records,
        'db_size': db_size,
        'db_size_str': f'{db_size / 1024:.2f} KB'
    })

# ========== 设备追踪（Android App上报） ==========
def init_device_tables():
    conn = get_db()
    conn.executescript('''
        CREATE TABLE IF NOT EXISTS user_devices (
            device_id TEXT PRIMARY KEY,
            nickname TEXT,
            first_seen TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            last_seen TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        );
        CREATE TABLE IF NOT EXISTS user_events (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            device_id TEXT,
            event_type TEXT,
            event_data TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        );
    ''')
    conn.commit()
    conn.close()

init_device_tables()

@app.route('/api/report', methods=['POST'])
def report_event():
    data = request.json or {}
    device_id = data.get('device_id', 'unknown')
    event_type = data.get('event_type', 'unknown')
    event_data = data.get('data', '')
    nickname = data.get('nickname', '')
    conn = get_db()
    conn.execute('INSERT INTO user_events (device_id, event_type, event_data) VALUES (?, ?, ?)',
                 (device_id, event_type, str(event_data)))
    conn.execute('''INSERT INTO user_devices (device_id, nickname, last_seen)
                    VALUES (?, ?, CURRENT_TIMESTAMP)
                    ON CONFLICT(device_id) DO UPDATE SET
                    nickname=COALESCE(NULLIF(excluded.nickname,''), nickname),
                    last_seen=CURRENT_TIMESTAMP''',
                 (device_id, nickname))
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

@app.route('/api/heartbeat', methods=['POST'])
def heartbeat():
    data = request.json or {}
    device_id = data.get('device_id', 'unknown')
    nickname = data.get('nickname', '')
    conn = get_db()
    conn.execute('''INSERT INTO user_devices (device_id, nickname, last_seen)
                    VALUES (?, ?, CURRENT_TIMESTAMP)
                    ON CONFLICT(device_id) DO UPDATE SET
                    nickname=COALESCE(NULLIF(excluded.nickname,''), nickname),
                    last_seen=CURRENT_TIMESTAMP''',
                 (device_id, nickname))
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

# ========== 管理面板 ==========
@app.route('/admin/login', methods=['POST'])
def admin_login():
    data = request.json or {}
    password = data.get('password', '')
    if bcrypt.checkpw(password.encode('utf-8'), ADMIN_PASSWORD_HASH.encode('utf-8')):
        session['admin_logged_in'] = True
        return jsonify({'ok': True})
    return jsonify({'error': '密码错误'}), 401

@app.route('/admin/logout', methods=['POST'])
def admin_logout():
    session.pop('admin_logged_in', None)
    return jsonify({'ok': True})

@app.route('/admin')
def admin_dashboard():
    return render_template('dashboard.html')

@app.route('/web/')
@app.route('/web')
def web_app():
    return send_from_directory(os.path.join(app.static_folder, 'web'), 'index.html')

@app.route('/api/admin/stats')
@admin_required
def admin_stats():
    conn = get_db()
    # 用户统计
    total_users = conn.execute('SELECT COUNT(*) FROM user_devices').fetchone()[0]
    active_today = conn.execute("SELECT COUNT(*) FROM user_devices WHERE date(last_seen)=date('now')").fetchone()[0]
    server_users = [dict(r) for r in conn.execute(
        'SELECT username, nickname, gender, banned, created_at, user_type, verify_status, verify_data FROM server_users ORDER BY created_at DESC'
    ).fetchall()]
    users = [dict(r) for r in conn.execute(
        'SELECT device_id, nickname, first_seen, last_seen FROM user_devices ORDER BY last_seen DESC'
    ).fetchall()]

    # 筑爱故事汇 (source=story)
    story_posts = [dict(r) for r in conn.execute(
        '''SELECT p.id, p.content, p.anonymous_name, p.device_id, p.media_url, p.tags, p.source,
           strftime('%s', p.created_at) as created_at,
           (SELECT COUNT(*) FROM anonymous_replies r WHERE r.post_id=p.id) as reply_count,
           (SELECT COUNT(*) FROM anonymous_likes l WHERE l.post_id=p.id) as like_count
           FROM anonymous_posts p WHERE p.source='story' ORDER BY p.created_at DESC LIMIT 100'''
    ).fetchall()]
    total_stories = conn.execute("SELECT COUNT(*) FROM anonymous_posts WHERE source='story'").fetchone()[0]

    # 叙心树洞 (source=treehole)
    treehole_posts = [dict(r) for r in conn.execute(
        '''SELECT p.id, p.content, p.anonymous_name, p.device_id, p.media_url, p.tags, p.source,
           strftime('%s', p.created_at) as created_at,
           (SELECT COUNT(*) FROM anonymous_replies r WHERE r.post_id=p.id) as reply_count,
           (SELECT COUNT(*) FROM anonymous_likes l WHERE l.post_id=p.id) as like_count
           FROM anonymous_posts p WHERE p.source='treehole' OR p.source IS NULL ORDER BY p.created_at DESC LIMIT 100'''
    ).fetchall()]
    total_treehole = conn.execute("SELECT COUNT(*) FROM anonymous_posts WHERE source='treehole' OR source IS NULL").fetchone()[0]

    # 投递树洞
    treehole_cases = []
    try:
        treehole_cases = [dict(r) for r in conn.execute(
            '''SELECT c.id, c.content, c.author_username, c.author_type, c.tags, c.status,
               c.like_count, c.created_at,
               (SELECT COUNT(*) FROM treehole_answers a WHERE a.case_id=c.id) as answer_count
               FROM treehole_cases c ORDER BY c.created_at DESC LIMIT 100'''
        ).fetchall()]
    except:
        pass
    total_cases = 0
    try:
        total_cases = conn.execute('SELECT COUNT(*) FROM treehole_cases').fetchone()[0]
    except:
        pass

    # 心晴评测
    mood_stats = [dict(r) for r in conn.execute(
        'SELECT mood, COUNT(*) as count FROM moods GROUP BY mood ORDER BY count DESC'
    ).fetchall()]
    total_moods = conn.execute('SELECT COUNT(*) FROM moods').fetchone()[0]

    # 测评结果
    test_results = []
    try:
        test_results = [dict(r) for r in conn.execute(
            'SELECT id, test_title, score, level, created_at FROM test_results_sync ORDER BY created_at DESC LIMIT 50'
        ).fetchall()]
    except:
        pass

    # AI对话
    total_chats = conn.execute('SELECT COUNT(*) FROM chat_history').fetchone()[0]
    chat_records = [dict(r) for r in conn.execute(
        'SELECT id, role, message, gender, created_at FROM chat_history ORDER BY created_at DESC LIMIT 100'
    ).fetchall()]

    # 回复统计
    total_replies = conn.execute('SELECT COUNT(*) FROM anonymous_replies').fetchone()[0]

    conn.close()
    return jsonify({
        'total_users': total_users, 'active_today': active_today,
        'server_users': server_users, 'users': users,
        'story_posts': story_posts, 'total_stories': total_stories,
        'treehole_posts': treehole_posts, 'total_treehole': total_treehole,
        'treehole_cases': treehole_cases, 'total_cases': total_cases,
        'mood_stats': mood_stats, 'total_moods': total_moods,
        'test_results': test_results,
        'total_chats': total_chats, 'chat_records': chat_records,
        'total_replies': total_replies
    })

# ========== 情感互助 ==========
def init_support_tables():
    conn = get_db()
    conn.executescript('''
        CREATE TABLE IF NOT EXISTS support_profiles (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            device_id TEXT UNIQUE NOT NULL,
            anonymous_name TEXT,
            mood_tags TEXT DEFAULT '[]',
            interest_tags TEXT DEFAULT '[]',
            bio TEXT DEFAULT '',
            is_active INTEGER DEFAULT 1,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        );
        CREATE TABLE IF NOT EXISTS support_matches (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user1_device_id TEXT NOT NULL,
            user2_device_id TEXT NOT NULL,
            status TEXT DEFAULT 'matched',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            UNIQUE(user1_device_id, user2_device_id)
        );
        CREATE TABLE IF NOT EXISTS support_messages (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            match_id INTEGER NOT NULL,
            sender_device_id TEXT NOT NULL,
            content TEXT NOT NULL,
            is_read INTEGER DEFAULT 0,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        );
    ''')
    # 兼容旧表：添加is_read字段
    try:
        conn.execute('ALTER TABLE support_messages ADD COLUMN is_read INTEGER DEFAULT 0')
    except:
        pass
    conn.commit()
    conn.close()

init_support_tables()

@app.route('/api/support/profile', methods=['POST'])
@limiter.limit("10 per minute")
def create_support_profile():
    data = request.json or {}
    device_id = data.get('device_id', '')
    mood_tags = json.dumps(data.get('mood_tags', []), ensure_ascii=False)
    interest_tags = json.dumps(data.get('interest_tags', []), ensure_ascii=False)
    bio = data.get('bio', '')
    if not device_id:
        return jsonify({'error': '设备ID不能为空'}), 400
    anon_name = get_anonymous_name(device_id)
    conn = get_db()
    conn.execute('''INSERT INTO support_profiles (device_id, anonymous_name, mood_tags, interest_tags, bio)
                    VALUES (?, ?, ?, ?, ?)
                    ON CONFLICT(device_id) DO UPDATE SET
                    mood_tags=excluded.mood_tags, interest_tags=excluded.interest_tags,
                    bio=excluded.bio, is_active=1, updated_at=CURRENT_TIMESTAMP''',
                 (device_id, anon_name, mood_tags, interest_tags, bio))
    conn.commit()
    conn.close()
    return jsonify({'ok': True, 'anonymous_name': anon_name})

@app.route('/api/support/profile', methods=['GET'])
def get_support_profile():
    device_id = request.args.get('device_id', '')
    if not device_id:
        return jsonify({'error': '设备ID不能为空'}), 400
    conn = get_db()
    profile = conn.execute('SELECT * FROM support_profiles WHERE device_id=?', (device_id,)).fetchone()
    conn.close()
    if not profile:
        return jsonify({'error': '名片不存在'}), 404
    d = dict(profile)
    d['mood_tags'] = json.loads(d['mood_tags']) if d['mood_tags'] else []
    d['interest_tags'] = json.loads(d['interest_tags']) if d['interest_tags'] else []
    return jsonify(d)

@app.route('/api/support/discover', methods=['GET'])
def discover_support_users():
    device_id = request.args.get('device_id', '')
    conn = get_db()
    # 获取自己的标签用于匹配
    my_profile = conn.execute('SELECT mood_tags, interest_tags FROM support_profiles WHERE device_id=?', (device_id,)).fetchone()
    my_moods = set(json.loads(my_profile['mood_tags'])) if my_profile and my_profile['mood_tags'] else set()
    my_interests = set(json.loads(my_profile['interest_tags'])) if my_profile and my_profile['interest_tags'] else set()

    # 获取已匹配的用户ID
    matched = conn.execute(
        'SELECT user1_device_id, user2_device_id FROM support_matches WHERE (user1_device_id=? OR user2_device_id=?) AND status="matched"',
        (device_id, device_id)
    ).fetchall()
    matched_ids = set()
    for m in matched:
        matched_ids.add(m['user1_device_id'])
        matched_ids.add(m['user2_device_id'])
    matched_ids.discard(device_id)
    # 获取活跃用户（排除自己和已匹配的）
    exclude_ids = matched_ids | {device_id}
    placeholders = ','.join('?' * len(exclude_ids))
    query = f'SELECT * FROM support_profiles WHERE is_active=1 AND device_id NOT IN ({placeholders})'
    params = list(exclude_ids)
    users = conn.execute(query, params).fetchall()
    # 计算匹配分数并获取真实昵称
    result = []
    for u in users:
        d = dict(u)
        d['mood_tags'] = json.loads(d['mood_tags']) if d['mood_tags'] else []
        d['interest_tags'] = json.loads(d['interest_tags']) if d['interest_tags'] else []
        # 匹配算法：情绪标签重合*10 + 兴趣标签重合*5
        mood_match = len(set(d['mood_tags']) & my_moods) * 10
        interest_match = len(set(d['interest_tags']) & my_interests) * 5
        d['match_score'] = mood_match + interest_match
        # 关联user_devices获取真实昵称，没有则查server_users
        ud = conn.execute('SELECT nickname FROM user_devices WHERE device_id=?', (d['device_id'],)).fetchone()
        if ud and ud['nickname']:
            d['real_name'] = ud['nickname']
        else:
            su = conn.execute('SELECT username, nickname FROM server_users WHERE username=?', (d['device_id'],)).fetchone()
            d['real_name'] = (su['nickname'] or su['username']) if su else ''
        result.append(d)
    # 按匹配分数降序排列
    result.sort(key=lambda x: x['match_score'], reverse=True)

    conn.close()
    return jsonify({'users': result[:20]})

@app.route('/api/support/match', methods=['POST'])
@limiter.limit("20 per minute")
def create_support_match():
    data = request.json or {}
    device_id = data.get('device_id', '')
    target_device_id = data.get('target_device_id', '')
    if not device_id or not target_device_id:
        return jsonify({'error': '参数缺失'}), 400
    if device_id == target_device_id:
        return jsonify({'error': '不能匹配自己'}), 400
    conn = get_db()
    # 检查是否已匹配
    existing = conn.execute(
        'SELECT id FROM support_matches WHERE ((user1_device_id=? AND user2_device_id=?) OR (user1_device_id=? AND user2_device_id=?))',
        (device_id, target_device_id, target_device_id, device_id)
    ).fetchone()
    if existing:
        conn.close()
        return jsonify({'error': '已经匹配过了'}), 400
    # 如果目标用户没有互助名片，自动创建
    target_profile = conn.execute('SELECT id FROM support_profiles WHERE device_id=?', (target_device_id,)).fetchone()
    if not target_profile:
        anon_name = get_anonymous_name(target_device_id)
        conn.execute('INSERT INTO support_profiles (device_id, anonymous_name, mood_tags, interest_tags, bio) VALUES (?,?,"[]","[]","")',
                     (target_device_id, anon_name))
    conn.execute('INSERT INTO support_matches (user1_device_id, user2_device_id) VALUES (?,?)',
                 (device_id, target_device_id))
    conn.commit()
    match_id = conn.execute('SELECT last_insert_rowid()').fetchone()[0]
    conn.close()
    return jsonify({'ok': True, 'match_id': match_id})

@app.route('/api/support/matches', methods=['GET'])
def get_support_matches():
    device_id = request.args.get('device_id', '')
    if not device_id:
        return jsonify({'matches': []})
    conn = get_db()
    matches = conn.execute(
        'SELECT * FROM support_matches WHERE (user1_device_id=? OR user2_device_id=?) AND status="matched" ORDER BY created_at DESC',
        (device_id, device_id)
    ).fetchall()
    result = []
    for m in matches:
        partner_id = m['user2_device_id'] if m['user1_device_id'] == device_id else m['user1_device_id']
        # 直接用真实用户名，先查设备表再查注册表
        ud = conn.execute('SELECT nickname FROM user_devices WHERE device_id=?', (partner_id,)).fetchone()
        if ud and ud['nickname']:
            partner_name = ud['nickname']
        else:
            su = conn.execute('SELECT username, nickname FROM server_users WHERE username=?', (partner_id,)).fetchone()
            partner_name = (su['nickname'] or su['username']) if su else '用户'
        # 获取最后一条消息
        last_msg = conn.execute(
            'SELECT content, created_at FROM support_messages WHERE match_id=? ORDER BY id DESC LIMIT 1',
            (m['id'],)
        ).fetchone()
        result.append({
            'match_id': m['id'],
            'partner_name': partner_name,
            'partner_device_id': partner_id,
            'last_message': last_msg['content'] if last_msg else '',
            'last_time': last_msg['created_at'] if last_msg else m['created_at'],
            'created_at': m['created_at']
        })
    conn.close()
    return jsonify({'matches': result})

@app.route('/api/support/messages', methods=['GET'])
def get_support_messages():
    match_id = request.args.get('match_id', '')
    if not match_id:
        return jsonify({'messages': []})
    conn = get_db()
    messages = conn.execute(
        "SELECT id, match_id, sender_device_id, content, is_read, strftime('%s', created_at) as created_at FROM support_messages WHERE match_id=? ORDER BY id ASC LIMIT 100",
        (match_id,)
    ).fetchall()
    conn.close()
    return jsonify({'messages': [dict(m) for m in messages]})

@app.route('/api/support/messages/read', methods=['POST'])
@limiter.limit("30 per minute")
def mark_messages_read():
    data = request.json or {}
    match_id = data.get('match_id', '')
    device_id = data.get('device_id', '')
    if not match_id or not device_id:
        return jsonify({'error': '参数缺失'}), 400
    conn = get_db()
    # 标记对方发来的消息为已读
    conn.execute(
        'UPDATE support_messages SET is_read=1 WHERE match_id=? AND sender_device_id!=? AND is_read=0',
        (match_id, device_id)
    )
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

@app.route('/api/support/messages', methods=['POST'])
@limiter.limit("30 per minute")
def send_support_message():
    data = request.json or {}
    match_id = data.get('match_id', '')
    device_id = data.get('device_id', '')
    content = data.get('content', '').strip()
    if not match_id or not device_id or not content:
        return jsonify({'error': '参数缺失'}), 400
    conn = get_db()
    # 验证是否是匹配成员
    match = conn.execute(
        'SELECT * FROM support_matches WHERE id=? AND (user1_device_id=? OR user2_device_id=?) AND status="matched"',
        (match_id, device_id, device_id)
    ).fetchone()
    if not match:
        conn.close()
        return jsonify({'error': '无权发送消息'}), 403
    conn.execute('INSERT INTO support_messages (match_id, sender_device_id, content) VALUES (?,?,?)',
                 (match_id, device_id, content))
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

@app.route('/api/support/search', methods=['GET'])
def search_support_users():
    device_id = request.args.get('device_id', '')
    keyword = request.args.get('keyword', '').strip()
    if not keyword:
        return jsonify({'users': []})
    conn = get_db()
    kw = f'%{keyword}%'
    # 1) 从 user_devices 搜索（APP注册用户），排除被封禁的
    users = conn.execute('''
        SELECT ud.device_id, ud.nickname as real_name,
               su.username as reg_username,
               COALESCE(sp.mood_tags, '[]') as mood_tags,
               COALESCE(sp.interest_tags, '[]') as interest_tags,
               COALESCE(sp.bio, '') as bio,
               CASE WHEN sp.device_id IS NOT NULL THEN 1 ELSE 0 END as has_profile
        FROM user_devices ud
        LEFT JOIN support_profiles sp ON ud.device_id = sp.device_id
        LEFT JOIN server_users su ON ud.nickname = su.nickname
        WHERE ud.device_id != ? AND (sp.device_id IS NULL OR sp.is_active=1) AND (ud.nickname LIKE ? OR su.username LIKE ?)
        LIMIT 20
    ''', (device_id, kw, kw)).fetchall()
    found_ids = {u['device_id'] for u in users}
    # 2) 从 server_users 搜索（网页注册用户），排除已有设备的和自己
    web_users = conn.execute('''
        SELECT su.username, su.nickname, su.gender
        FROM server_users su
        LEFT JOIN support_profiles sp ON su.username = sp.device_id
        WHERE (sp.device_id IS NULL OR sp.is_active=1) AND su.username != ? AND (su.username LIKE ? OR su.nickname LIKE ?)
        LIMIT 20
    ''', (device_id, kw, kw)).fetchall()
    conn.close()
    result = []
    for u in users:
        d = dict(u)
        d['mood_tags'] = json.loads(d['mood_tags']) if d['mood_tags'] else []
        d['interest_tags'] = json.loads(d['interest_tags']) if d['interest_tags'] else []
        real = d.get('reg_username') or d.get('real_name') or '用户'
        d['real_name'] = real
        result.append(d)
    return jsonify({'users': result})

@app.route('/api/support/unmatch', methods=['POST'])
@limiter.limit("10 per minute")
def unmatch_support():
    data = request.json or {}
    device_id = data.get('device_id', '')
    match_id = data.get('match_id', '')
    if not device_id or not match_id:
        return jsonify({'error': '参数缺失'}), 400
    conn = get_db()
    # 验证是否是匹配成员
    match = conn.execute(
        'SELECT * FROM support_matches WHERE id=? AND (user1_device_id=? OR user2_device_id=?)',
        (match_id, device_id, device_id)
    ).fetchone()
    if not match:
        conn.close()
        return jsonify({'error': '匹配不存在'}), 404
    # 删除匹配记录
    conn.execute('DELETE FROM support_matches WHERE id=?', (match_id,))
    # 删除相关聊天记录
    conn.execute('DELETE FROM support_messages WHERE match_id=?', (match_id,))
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

@app.route('/api/support/report', methods=['POST'])
@limiter.limit("5 per minute")
def report_support_user():
    data = request.json or {}
    device_id = data.get('device_id', '')
    target_device_id = data.get('target_device_id', '')
    reason = data.get('reason', '')
    if not device_id or not target_device_id:
        return jsonify({'error': '参数缺失'}), 400
    conn = get_db()
    # 封禁被举报用户
    conn.execute('UPDATE support_profiles SET is_active=0 WHERE device_id=?', (target_device_id,))
    # 解除匹配关系
    conn.execute("UPDATE support_matches SET status='blocked' WHERE (user1_device_id=? OR user2_device_id=?) AND status='matched'",
                 (target_device_id, target_device_id))
    conn.commit()
    conn.close()
    return jsonify({'ok': True})

# ========== 志愿者申请 ==========
def init_volunteer_table():
    conn = get_db()
    conn.execute('''CREATE TABLE IF NOT EXISTS volunteers (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT NOT NULL,
        reason TEXT DEFAULT '',
        skills TEXT DEFAULT '',
        status TEXT DEFAULT 'pending',
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')
    conn.commit()
    conn.close()

init_volunteer_table()

@app.route('/api/volunteer/apply', methods=['POST'])
@limiter.limit("5 per minute")
def apply_volunteer():
    data = request.json or {}
    username = data.get('username', '')
    reason = data.get('reason', '')
    skills = data.get('skills', '')
    if not username:
        return jsonify({'error': '请先登录'}), 400
    conn = get_db()
    existing = conn.execute('SELECT id FROM volunteers WHERE username=? AND status="pending"', (username,)).fetchone()
    if existing:
        conn.close()
        return jsonify({'error': '您已有待审核的申请'}), 400
    conn.execute('INSERT INTO volunteers (username, reason, skills) VALUES (?,?,?)', (username, reason, skills))
    conn.commit()
    conn.close()
    return jsonify({'ok': True, 'message': '申请已提交，等待审核'})

@app.route('/api/volunteer/status', methods=['GET'])
def volunteer_status():
    username = request.args.get('username', '')
    if not username:
        return jsonify({'status': 'none'})
    conn = get_db()
    row = conn.execute('SELECT status FROM volunteers WHERE username=? ORDER BY created_at DESC LIMIT 1', (username,)).fetchone()
    conn.close()
    return jsonify({'status': row['status'] if row else 'none'})

# ========== AI情绪评估 ==========
@app.route('/api/ai-emotion-assess', methods=['POST'])
@limiter.limit("10 per minute")
def ai_emotion_assess():
    data = request.get_json()
    description = data.get('description', '').strip()
    device_id = data.get('device_id', '')
    if not description:
        return jsonify({'error': '请描述您的情绪状态'}), 400

    system_prompt = """你是一位专业的心理咨询师和情绪评估专家。请根据用户的情绪描述进行专业评估。

请从以下几个维度进行分析：
1. 情绪识别：识别用户当前的主要情绪（如焦虑、抑郁、愤怒、恐惧、悲伤等）
2. 情绪强度：评估情绪的严重程度（轻度/中度/重度）
3. 影响分析：分析该情绪可能对生活、学习、社交的影响
4. 风险等级：低风险/中风险/高风险
5. 应对建议：给出具体可行的调节方法
6. 专业建议：如果情况严重，建议寻求专业帮助

请用温暖、专业的语气回复，不要使用过于专业的术语，让用户容易理解。回复格式为结构化的评估报告。"""

    messages = [{"role": "user", "content": f"请对以下情绪描述进行评估：\n\n{description}"}]

    try:
        api_url = "https://token-plan-cn.xiaomimimo.com/anthropic/v1/messages"
        headers = {
            "Content-Type": "application/json",
            "x-api-key": os.environ.get('MIMO_API_KEY', ''),
            "anthropic-version": "2023-06-01"
        }
        payload = {
            "model": "mimo-v2.5",
            "system": system_prompt,
            "max_tokens": 2048,
            "messages": messages
        }
        resp = requests.post(api_url, headers=headers, json=payload, timeout=30)
        result = resp.json()
        content_list = result.get('content', [])
        assessment = ''
        for item in content_list:
            if item.get('type') == 'text':
                assessment = item.get('text', '')
                break
            elif item.get('type') == 'thinking':
                assessment = item.get('thinking', '')
        if not assessment:
            assessment = '抱歉，暂时无法完成评估，请稍后再试。'
    except Exception as e:
        assessment = "抱歉，AI情绪评估暂时不可用。如果您需要帮助，请拨打心理援助热线：12320-5"

    return jsonify({'assessment': assessment, 'description': description})

# ========== 活动和课程预约 ==========
@app.route('/api/events', methods=['GET'])
def get_events():
    events = [
        {"id": 1, "title": "校园心理讲座", "desc": "主题：压力管理与情绪调节", "time": "2026-05-20 14:00", "location": "学术报告厅A", "capacity": 100, "booked": 45, "tags": ["心理", "讲座"]},
        {"id": 2, "title": "户外拓展活动", "desc": "团队协作与信任建立", "time": "2026-05-25 09:00", "location": "操场", "capacity": 50, "booked": 30, "tags": ["户外", "团队"]},
        {"id": 3, "title": "读书分享会", "desc": "分享《被讨厌的勇气》读后感", "time": "2026-05-22 18:30", "location": "图书馆3楼", "capacity": 30, "booked": 12, "tags": ["读书", "分享"]},
        {"id": 4, "title": "正念冥想体验", "desc": "学习正念冥想，缓解焦虑", "time": "2026-05-28 16:00", "location": "心理健康中心", "capacity": 20, "booked": 18, "tags": ["冥想", "减压"]},
        {"id": 5, "title": "艺术疗愈工作坊", "desc": "通过绘画表达内心情感", "time": "2026-06-01 14:00", "location": "艺术教室", "capacity": 25, "booked": 8, "tags": ["艺术", "疗愈"]},
    ]
    return jsonify({'events': events})

@app.route('/api/classes', methods=['GET'])
def get_classes():
    classes = [
        {"id": 1, "title": "情绪管理基础课", "desc": "6节课学会识别和管理情绪", "schedule": "每周三 19:00", "duration": "6周", "capacity": 40, "booked": 28, "tags": ["情绪", "基础"]},
        {"id": 2, "title": "人际沟通技巧", "desc": "提升沟通能力，改善人际关系", "schedule": "每周四 19:00", "duration": "4周", "capacity": 35, "booked": 15, "tags": ["沟通", "人际"]},
        {"id": 3, "title": "考试焦虑应对", "desc": "考前减压，提升应试心理素质", "schedule": "每周二 20:00", "duration": "3周", "capacity": 50, "booked": 42, "tags": ["考试", "减压"]},
        {"id": 4, "title": "自我认知探索", "desc": "深入了解自己，找到人生方向", "schedule": "每周五 18:30", "duration": "5周", "capacity": 30, "booked": 10, "tags": ["自我", "成长"]},
        {"id": 5, "title": "睡眠质量改善", "desc": "科学方法改善睡眠", "schedule": "每周一 20:00", "duration": "4周", "capacity": 45, "booked": 20, "tags": ["睡眠", "健康"]},
    ]
    return jsonify({'classes': classes})

@app.route('/api/events/<int:event_id>/book', methods=['POST'])
@limiter.limit("10 per minute")
def book_event(event_id):
    data = request.json or {}
    device_id = data.get('device_id', '')
    if not device_id:
        return jsonify({'error': '请先登录'}), 400
    return jsonify({'ok': True, 'message': '预约成功'})

@app.route('/api/classes/<int:class_id>/book', methods=['POST'])
@limiter.limit("10 per minute")
def book_class(class_id):
    data = request.json or {}
    device_id = data.get('device_id', '')
    if not device_id:
        return jsonify({'error': '请先登录'}), 400
    return jsonify({'ok': True, 'message': '预约成功'})

if __name__ == '__main__':
    app.run(debug=True, port=5000)
